#!/usr/bin/env python3
"""Génère le rapport Word complet pour MicrofinanceHub."""

import io
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.patheffects as pe
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Couleurs thème ──────────────────────────────────────────────────────────
VERT     = (0.024, 0.306, 0.243)
VERT_L   = (0.064, 0.618, 0.486)
BLEU     = (0.09,  0.44,  0.73)
ORANGE   = (0.85,  0.40,  0.10)
GRIS     = (0.94,  0.94,  0.94)
BLANC    = (1.0,   1.0,   1.0)
ROUGE    = (0.72,  0.11,  0.11)

def fig_to_stream(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)
    return buf

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMME 1 : Architecture Globale Microservices
# ══════════════════════════════════════════════════════════════════════════════
def draw_architecture_globale():
    fig, ax = plt.subplots(figsize=(16, 10))
    ax.set_xlim(0, 16); ax.set_ylim(0, 10); ax.axis('off')
    fig.patch.set_facecolor('#f8f9fa')

    def box(x, y, w, h, label, sub='', color=VERT, lc=BLANC, fc=BLANC, fs=8):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.05",
                     linewidth=1.5, edgecolor=color, facecolor=fc))
        ax.text(x+w/2, y+h/2 + (0.12 if sub else 0), label,
                ha='center', va='center', fontsize=fs, fontweight='bold', color=color)
        if sub:
            ax.text(x+w/2, y+h/2 - 0.15, sub,
                    ha='center', va='center', fontsize=6, color='#666666')

    def cloud(x, y, w, h, label, color=BLEU):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.1",
                     linewidth=2, edgecolor=color, facecolor=(*color[:3], 0.12), linestyle='--'))
        ax.text(x+w/2, y+h+0.12, label, ha='center', va='bottom',
                fontsize=8, fontweight='bold', color=color)

    def arrow(x1, y1, x2, y2, col='#555555'):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle='->', color=col, lw=1.5))

    cloud(0.1, 1.0, 2.4, 7.5, 'CLIENTS', color=(0.4, 0.4, 0.4))
    box(0.3, 7.2, 2.0, 0.7, 'Navigateur', 'React SPA', (0.4,0.4,0.4))
    box(0.3, 6.2, 2.0, 0.7, 'Application', 'Mobile / Tablette', (0.4,0.4,0.4))
    box(0.3, 5.2, 2.0, 0.7, 'Systemes tiers', 'CamPay / Twilio', (0.4,0.4,0.4))

    box(2.9, 6.4, 2.2, 1.2, 'Frontend React', ':3000', VERT_L, fc=(*VERT_L, 0.06))
    box(5.4, 6.4, 2.2, 1.2, 'API Gateway', ':8091', BLEU, fc=(*BLEU, 0.06), fs=8)

    cloud(5.2, 4.5, 2.6, 1.6, 'Spring Cloud', color=ORANGE)
    box(5.4, 4.7, 1.1, 1.1, 'Config', ':8000', ORANGE, fc=(*ORANGE, 0.08), fs=7)
    box(6.6, 4.7, 1.1, 1.1, 'Eureka', ':8761', ORANGE, fc=(*ORANGE, 0.08), fs=7)

    cloud(8.2, 1.0, 7.6, 8.5, 'Microservices Metier', color=VERT)

    services = [
        ('Auth',        ':8080', 8.4,  8.0),
        ('Client',      ':8081', 9.8,  8.0),
        ('Account',     ':8082', 11.2, 8.0),
        ('Loan',        ':8083', 12.6, 8.0),
        ('Repayment',   ':8084', 14.0, 8.0),
        ('Reporting',   ':8085', 8.4,  6.0),
        ('Agency',      ':8086', 9.8,  6.0),
        ('Config.MFI',  ':8087', 11.2, 6.0),
        ('Transaction', ':8088', 12.6, 6.0),
        ('Notif.',      ':8089', 14.0, 6.0),
    ]
    for label, port, sx, sy in services:
        box(sx, sy, 1.3, 0.85, label, port, VERT, fc=(*VERT, 0.05), fs=6.5)

    cloud(8.2, 2.0, 7.6, 3.5, 'Infrastructure de Donnees', color=(0.5, 0.2, 0.6))
    box(8.4,  3.2, 1.5, 1.0, 'PostgreSQL', ':5433', (0.5,0.2,0.6), fc=(0.5,0.2,0.6,0.06), fs=7)
    box(10.1, 3.2, 1.5, 1.0, 'RabbitMQ', ':5672', (0.5,0.2,0.6), fc=(0.5,0.2,0.6,0.06), fs=7)
    box(11.8, 3.2, 1.5, 1.0, 'Redis', ':6379', (0.5,0.2,0.6), fc=(0.5,0.2,0.6,0.06), fs=7)
    box(13.5, 3.2, 1.5, 1.0, 'Google\nDrive', 'API', (0.5,0.2,0.6), fc=(0.5,0.2,0.6,0.06), fs=7)

    arrow(2.5, 7.0, 2.9, 7.0); arrow(2.5, 6.0, 2.9, 6.5)
    arrow(5.1, 7.0, 5.4, 7.0)
    arrow(7.6, 7.0, 8.2, 8.2)
    arrow(6.5, 6.4, 6.8, 6.2, ORANGE); arrow(6.8, 6.2, 8.2, 7.5)
    for sx, sy in [(s[2]+0.65, s[3]) for s in services]:
        ax.plot([sx], [sy-0.85], 'v', color=VERT, markersize=4, alpha=0.5)

    ax.set_title('Architecture Globale — MicrofinanceHub',
                 fontsize=14, fontweight='bold', color='#064e3b', pad=12)
    return fig_to_stream(fig)

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMME 2 : Architecture Technique Stack
# ══════════════════════════════════════════════════════════════════════════════
def draw_stack_technique():
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14); ax.set_ylim(0, 8); ax.axis('off')
    fig.patch.set_facecolor('#f8f9fa')

    cols = [
        ('FRONTEND',       1.0,  ['React 18', 'Axios / REST', 'Tailwind CSS', 'React Router', 'Lucide Icons', 'JWT (localStorage)'], VERT_L),
        ('BACKEND',        4.5,  ['Spring Boot 3.4.4', 'Spring Cloud 2024', 'Spring Security', 'Spring Data JPA', 'OpenFeign', 'JJWT 0.12.6'], BLEU),
        ('MESSAGERIE',     8.0,  ['RabbitMQ 3.13', 'AMQP Protocol', 'Echanges direct/fanout', 'Consumers async', 'Dead Letter Queue', 'Management UI'], (0.7, 0.3, 0.0)),
        ('INFRA / DEVOPS', 11.4, ['Docker + Compose', 'PostgreSQL 16', 'Redis 7', 'Eureka Discovery', 'Config Server', 'GitHub'], (0.4, 0.2, 0.6)),
    ]

    for title, x, items, col in cols:
        ax.add_patch(FancyBboxPatch((x, 6.5), 2.8, 1.2, boxstyle="round,pad=0.08",
                     facecolor=col, edgecolor=col, linewidth=0))
        ax.text(x+1.4, 7.1, title, ha='center', va='center',
                fontsize=9, fontweight='bold', color='white')
        for i, item in enumerate(items):
            y = 5.6 - i * 0.85
            fc = (*col, 0.08) if i % 2 == 0 else BLANC
            ax.add_patch(FancyBboxPatch((x, y), 2.8, 0.75, boxstyle="round,pad=0.04",
                         facecolor=fc, edgecolor=(*col, 0.3), linewidth=0.8))
            ax.text(x+1.4, y+0.375, item, ha='center', va='center',
                    fontsize=8, color='#222222')

    ax.set_title('Stack Technique — MicrofinanceHub',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMME 3 : Diagramme de Cas d'Utilisation (Use Case)
# ══════════════════════════════════════════════════════════════════════════════
def draw_use_case():
    fig, ax = plt.subplots(figsize=(16, 11))
    ax.set_xlim(0, 16); ax.set_ylim(0, 11); ax.axis('off')
    fig.patch.set_facecolor('#f8f9fa')

    def actor(x, y, label, color='#333333'):
        ax.add_patch(plt.Circle((x, y+1.1), 0.25, color=color, zorder=5))
        ax.plot([x, x], [y+0.85, y+0.3], color=color, lw=2, zorder=5)
        ax.plot([x-0.35, x+0.35], [y+0.65, y+0.65], color=color, lw=2, zorder=5)
        ax.plot([x, x-0.3], [y+0.3, y-0.1], color=color, lw=2, zorder=5)
        ax.plot([x, x+0.3], [y+0.3, y-0.1], color=color, lw=2, zorder=5)
        ax.text(x, y-0.35, label, ha='center', va='top', fontsize=8,
                fontweight='bold', color=color)

    def uc(x, y, w, h, label):
        ax.add_patch(mpatches.Ellipse((x, y), w, h, color='white',
                     ec=VERT_L, lw=1.5, zorder=3))
        ax.text(x, y, label, ha='center', va='center', fontsize=7,
                color='#064e3b', zorder=4)

    def assoc(ax_x, ay, bx, by):
        ax.plot([ax_x, bx], [ay, by], 'k-', lw=0.8, zorder=2)

    ax.add_patch(FancyBboxPatch((3.0, 0.5), 12.5, 9.8, boxstyle="round,pad=0.1",
                 facecolor='#f0f4f8', edgecolor=BLEU, linewidth=2, zorder=0))
    ax.text(9.25, 10.5, 'Systeme MicrofinanceHub', ha='center', fontsize=10,
            fontweight='bold', color='#1a70ba', style='italic')

    actor(1.2, 8.5, 'Admin',      '#7c3aed')
    actor(1.2, 5.5, 'Directeur',  '#064e3b')
    actor(1.2, 2.5, 'Agent',      '#0369a1')
    actor(14.8,5.0, 'Client',     '#b45309')
    actor(14.8,2.0, 'CamPay\n(Ext.)', '#6b7280')

    cases_admin = [
        (5.5, 9.5, 3.2, 0.55, 'Gerer les utilisateurs'),
        (5.5, 8.7, 3.2, 0.55, 'Creer/Gerer les agences'),
        (5.5, 7.9, 3.2, 0.55, 'Assigner directeurs/agents'),
        (5.5, 7.1, 3.2, 0.55, 'Valider comptes & paiements'),
        (5.5, 6.3, 3.2, 0.55, 'Consulter rapports globaux'),
    ]
    for args in cases_admin:
        uc(*args)
        assoc(2.0, 8.8, args[0]-1.6, args[1])

    cases_dir = [
        (9.2, 9.5, 3.2, 0.55, 'Gerer clients (agence)'),
        (9.2, 8.7, 3.2, 0.55, 'Gerer agents (agence)'),
        (9.2, 7.9, 3.2, 0.55, 'Valider comptes (agence)'),
        (9.2, 7.1, 3.2, 0.55, 'Effectuer transactions'),
        (9.2, 6.3, 3.2, 0.55, 'Rapports agence'),
    ]
    for args in cases_dir:
        uc(*args)
        assoc(2.0, 6.0, args[0]-1.6, args[1])

    cases_agent = [
        (5.5, 5.3, 3.2, 0.55, 'Creer clients'),
        (5.5, 4.5, 3.2, 0.55, 'Ouvrir comptes'),
        (5.5, 3.7, 3.2, 0.55, 'Operations guichet'),
        (5.5, 2.9, 3.2, 0.55, 'Enregistrer paiements'),
        (5.5, 2.1, 3.2, 0.55, 'Demandes de pret'),
    ]
    for args in cases_agent:
        uc(*args)
        assoc(2.0, 3.0, args[0]-1.6, args[1])

    cases_client = [
        (9.2, 5.3, 3.2, 0.55, "S'inscrire / Se connecter"),
        (9.2, 4.5, 3.2, 0.55, 'Consulter compte & solde'),
        (9.2, 3.7, 3.2, 0.55, 'Effectuer paiements'),
        (9.2, 2.9, 3.2, 0.55, 'Suivi des prets'),
        (9.2, 2.1, 3.2, 0.55, 'Paiement Mobile Money'),
    ]
    for args in cases_client:
        uc(*args)
        assoc(14.2, 5.2, args[0]+1.6, args[1])

    assoc(14.2, 2.3, 10.8, 2.1)
    assoc(14.2, 2.3, 10.8, 2.9)

    ax.set_title("Diagramme de Cas d'Utilisation — MicrofinanceHub",
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMME 4 : Diagramme de Séquence — Authentification
# ══════════════════════════════════════════════════════════════════════════════
def draw_sequence_auth():
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14); ax.set_ylim(0, 8); ax.axis('off')
    fig.patch.set_facecolor('#f8f9fa')

    actors = [
        ('Utilisateur', 1.5,  (0.3,0.3,0.3)),
        ('Frontend\nReact', 3.8, VERT_L),
        ('Auth\nService', 6.5, (0.5,0.2,0.6)),
        ('PostgreSQL\n(Auth DB)', 9.5, (0.3,0.5,0.8)),
        ('Redis\n(Sessions)', 12.2, (0.7,0.1,0.1)),
    ]

    for label, x, col in actors:
        ax.add_patch(FancyBboxPatch((x-0.7, 7.0), 1.4, 0.8, boxstyle="round,pad=0.06",
                     facecolor=col, edgecolor=col))
        ax.text(x, 7.4, label, ha='center', va='center', fontsize=8,
                fontweight='bold', color='white')
        ax.plot([x, x], [0.5, 7.0], color='#bbbbbb', lw=1, linestyle='--')

    def msg(y, x1, x2, label, ret=False, color='#333333'):
        ax.annotate('', xy=(x2, y), xytext=(x1, y),
            arrowprops=dict(arrowstyle=('<-' if ret else '->'), color=color, lw=1.3))
        ax.text((x1+x2)/2, y+0.12, label, ha='center', va='bottom', fontsize=7.5, color=color)

    msg(6.5, 1.5, 3.8, 'Saisit email + mot de passe', color=(0.3,0.3,0.3))
    msg(6.0, 3.8, 6.5, 'POST /auth/login', color=VERT_L)
    msg(5.5, 6.5, 9.5, 'SELECT user WHERE email=?', color=(0.5,0.2,0.6))
    msg(5.0, 9.5, 6.5, 'UserDetails (hash, role)', ret=True, color=(0.3,0.5,0.8))
    msg(4.5, 6.5, 6.5, 'BCrypt.verify(password)', color=(0.5,0.2,0.6))
    ax.add_patch(FancyBboxPatch((6.0, 4.1), 1.0, 0.7, facecolor=(*VERT_L, 0.3),
                 edgecolor=VERT_L, linewidth=1))
    ax.text(6.5, 4.45, 'Genere\nJWT', ha='center', va='center', fontsize=7)
    msg(3.9, 6.5, 12.2, 'Stocke session (TTL)', color=(0.5,0.2,0.6))
    msg(3.4, 12.2, 6.5, 'Session creee', ret=True, color=(0.7,0.1,0.1))
    msg(2.9, 6.5, 3.8, '{ access_token, role, agencyId }', ret=True, color=(0.5,0.2,0.6))
    msg(2.4, 3.8, 1.5, 'Stocke JWT -> localStorage', ret=True, color=VERT_L)
    msg(1.9, 3.8, 1.5, 'Redirige vers /dashboard', ret=True, color=VERT_L)

    ax.set_title('Diagramme de Sequence — Authentification JWT',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMME 5 : Diagramme de Séquence — Processus de Prêt
# ══════════════════════════════════════════════════════════════════════════════
def draw_sequence_pret():
    fig, ax = plt.subplots(figsize=(16, 10))
    ax.set_xlim(0, 16); ax.set_ylim(0, 10); ax.axis('off')
    fig.patch.set_facecolor('#f8f9fa')

    actors = [
        ('Agent/\nDirecteur', 1.5,  BLEU),
        ('Frontend\nReact',   3.5,  VERT_L),
        ('API\nGateway',      5.5,  ORANGE),
        ('Auth\nService',     7.5,  (0.5,0.2,0.6)),
        ('Loan\nService',     9.5,  VERT),
        ('Client\nService',   11.5, (0.6,0.4,0.0)),
        ('Notif.\nService',   13.5, ROUGE),
    ]

    for label, x, col in actors:
        ax.add_patch(FancyBboxPatch((x-0.6, 8.9), 1.2, 0.9, boxstyle="round,pad=0.06",
                     facecolor=col, edgecolor=col, linewidth=0))
        ax.text(x, 9.35, label, ha='center', va='center',
                fontsize=7.5, fontweight='bold', color='white')
        ax.plot([x, x], [0.3, 8.9], color='#aaaaaa', lw=1, linestyle='--', zorder=1)

    def msg(y, x1, x2, label, ret=False, color='#333333', note=''):
        style = '->' if not ret else '<-'
        ax.annotate('', xy=(x2, y), xytext=(x1, y),
            arrowprops=dict(arrowstyle=style, color=color, lw=1.3))
        mx = (x1 + x2) / 2
        ax.text(mx, y + 0.12, label, ha='center', va='bottom', fontsize=7, color=color)
        if note:
            ax.text(mx, y - 0.15, f'({note})', ha='center', va='top',
                    fontsize=6, color='#888888', style='italic')

    def box_act(x, y, h=0.45, col=BLEU):
        ax.add_patch(FancyBboxPatch((x-0.12, y), 0.24, h,
                     boxstyle="square", facecolor=(*col, 0.3), edgecolor=col, lw=1))

    msg(8.4, 1.5, 3.5, '1. Saisie demande de pret', color=BLEU)
    msg(7.9, 3.5, 5.5, '2. POST /api/loans/apply', color=VERT_L)
    msg(7.4, 5.5, 7.5, '3. Valide JWT', color=ORANGE)
    msg(6.9, 7.5, 5.5, 'Token valide', ret=True, color=(0.5,0.2,0.6))
    msg(6.4, 5.5, 9.5, '4. Traite demande', color=ORANGE)
    msg(5.9, 9.5, 11.5, '5. Verifie eligibilite client', color=VERT)
    msg(5.4, 11.5, 9.5, 'Profil client OK', ret=True, color=(0.6,0.4,0.0))
    msg(4.9, 9.5, 9.5, '6. Calcule amortissement', color=VERT)
    box_act(9.5, 4.55, 0.7, VERT)
    msg(4.4, 9.5, 9.5, '7. Enregistre demande', color=VERT)
    msg(3.9, 9.5, 13.5, '8. Notif. async (RabbitMQ)', color=VERT)
    msg(3.4, 9.5, 5.5, 'Demande creee', ret=True, color=VERT)
    msg(2.9, 5.5, 3.5, 'Reponse applicant', ret=True, color=ORANGE)
    msg(2.4, 3.5, 1.5, "Confirmation a l'agent", ret=True, color=VERT_L)
    msg(1.9, 13.5, 1.5, '9. Email notification', color=ROUGE)

    ax.text(0.3, 0.9, 'RabbitMQ = messagerie asynchrone entre services',
            fontsize=7, style='italic', color='#666666')
    ax.text(0.3, 0.6, 'JWT = JSON Web Token pour authentification stateless',
            fontsize=7, style='italic', color='#666666')

    ax.set_title('Diagramme de Sequence — Demande de Pret',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMME 6 : Séquence — Dépôt d'argent (NOUVEAU)
# ══════════════════════════════════════════════════════════════════════════════
def draw_sequence_transaction():
    fig, ax = plt.subplots(figsize=(16, 10))
    ax.set_xlim(0, 16); ax.set_ylim(0, 10); ax.axis('off')
    fig.patch.set_facecolor('#f8f9fa')

    actors = [
        ('Agent/\nDirecteur', 1.2,  BLEU),
        ('Frontend\nReact',   3.2,  VERT_L),
        ('API\nGateway',      5.2,  ORANGE),
        ('Auth\nService',     7.2,  (0.5,0.2,0.6)),
        ('Transaction\nService', 9.5, VERT),
        ('Account\nService',  12.0, (0.3,0.5,0.8)),
        ('Notification\nService', 14.5, ROUGE),
    ]

    for label, x, col in actors:
        ax.add_patch(FancyBboxPatch((x-0.65, 8.9), 1.3, 0.9, boxstyle="round,pad=0.06",
                     facecolor=col, edgecolor=col, linewidth=0))
        ax.text(x, 9.35, label, ha='center', va='center',
                fontsize=7, fontweight='bold', color='white')
        ax.plot([x, x], [0.5, 8.9], color='#aaaaaa', lw=1, linestyle='--', zorder=1)

    def msg(y, x1, x2, label, ret=False, color='#333333', note=''):
        style = '->' if not ret else '<-'
        ax.annotate('', xy=(x2, y), xytext=(x1, y),
            arrowprops=dict(arrowstyle=style, color=color, lw=1.3))
        mx = (x1 + x2) / 2
        ax.text(mx, y + 0.13, label, ha='center', va='bottom', fontsize=6.8, color=color)
        if note:
            ax.text(mx, y - 0.12, f'({note})', ha='center', va='top',
                    fontsize=6, color='#888888', style='italic')

    def activation(x, y_start, y_end, col):
        ax.add_patch(FancyBboxPatch((x-0.1, y_end), 0.2, y_start - y_end,
                     boxstyle="square", facecolor=(*col, 0.25), edgecolor=col, lw=0.8))

    msg(8.4, 1.2, 3.2, '1. Saisie montant + compteId', color=BLEU)
    msg(7.9, 3.2, 5.2, '2. POST /api/transactions/depot/{compteId}', color=VERT_L)
    msg(7.4, 5.2, 7.2, '3. Validation JWT', color=ORANGE)
    msg(6.9, 7.2, 5.2, 'Token OK + roles extraits', ret=True, color=(0.5,0.2,0.6))
    msg(6.4, 5.2, 9.5, '4. Route vers Transaction Service', color=ORANGE)
    activation(9.5, 6.4, 4.0, VERT)
    msg(5.9, 9.5, 9.5, '5. Verifie compte actif + solde', color=VERT)
    msg(5.4, 9.5, 12.0, '6. POST /api/comptes/{id}/crediter (interne)', color=VERT)
    activation(12.0, 5.4, 4.5, (0.3,0.5,0.8))
    msg(4.9, 12.0, 12.0, '7. Maj solde en base PostgreSQL', color=(0.3,0.5,0.8))
    msg(4.4, 12.0, 9.5, 'Nouveau solde confirme', ret=True, color=(0.3,0.5,0.8))
    msg(3.9, 9.5, 9.5, '8. Enregistre transaction (type=DEPOT)', color=VERT)
    msg(3.4, 9.5, 14.5, '9. Publish event (RabbitMQ async)', color=VERT,
        note='transaction.completed')
    msg(2.9, 9.5, 5.2, 'Reponse: transaction creee', ret=True, color=VERT)
    msg(2.4, 5.2, 3.2, 'HTTP 201 Created', ret=True, color=ORANGE)
    msg(1.9, 3.2, 1.2, 'Affiche confirmation + nouveau solde', ret=True, color=VERT_L)
    msg(1.4, 14.5, 1.2, '10. Email/SMS notification client', color=ROUGE)

    ax.text(0.3, 1.1, 'Appel interne (Feign) entre Transaction Service et Account Service',
            fontsize=6.5, style='italic', color='#666666')
    ax.text(0.3, 0.8, 'Notification envoyee de maniere asynchrone via RabbitMQ',
            fontsize=6.5, style='italic', color='#666666')

    ax.set_title("Diagramme de Sequence — Depot d'Argent",
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)

# ══════════════════════════════════════════════════════════════════════════════
# HELPER UML  (partagé entre tous les diagrammes de classes)
# ══════════════════════════════════════════════════════════════════════════════
def _uml(ax, ROW_H=0.30, H_HDR=0.48):
    """Retourne cls(), enu(), arr() prêts à dessiner."""

    def cls(x, y, name, attrs, color, w=4.0):
        h_body = max(len(attrs) * ROW_H + 0.12, 0.35)
        total  = H_HDR + h_body
        ax.add_patch(FancyBboxPatch((x, y - total), w, total,
            boxstyle='square', facecolor='white', edgecolor=color, lw=1.4))
        ax.add_patch(FancyBboxPatch((x, y - H_HDR), w, H_HDR,
            boxstyle='square', facecolor=color, edgecolor=color, lw=1.4))
        ax.text(x + w/2, y - H_HDR/2, name,
                ha='center', va='center', fontsize=8.5, fontweight='bold', color='white')
        ax.plot([x, x+w], [y - H_HDR, y - H_HDR], color=color, lw=0.8)
        for i, attr in enumerate(attrs):
            ax.text(x + 0.12, y - H_HDR - 0.15 - i*ROW_H, attr,
                    va='center', fontsize=7, color='#333333', family='monospace')
        return y - total   # bottom y

    def enu(x, y, name, values, color, w=3.5):
        h_body = max(len(values) * 0.26 + 0.10, 0.28)
        total  = H_HDR + h_body
        ax.add_patch(FancyBboxPatch((x, y - total), w, total, boxstyle='square',
            facecolor='#fafafa', edgecolor=color, lw=1.0, linestyle='--'))
        ax.add_patch(FancyBboxPatch((x, y - H_HDR), w, H_HDR, boxstyle='square',
            facecolor=(*color, 0.65), edgecolor=color, lw=1.0))
        ax.text(x + w/2, y - H_HDR/2, f'«enumeration»\n{name}',
                ha='center', va='center', fontsize=7.2, fontweight='bold', color='white')
        ax.plot([x, x+w], [y - H_HDR, y - H_HDR], color=color, lw=0.7, linestyle='--')
        for i, v in enumerate(values):
            ax.text(x + w/2, y - H_HDR - 0.12 - i*0.26, v,
                    ha='center', va='center', fontsize=6.8, color='#444444')
        return y - total

    def arr(x1, y1, x2, y2, label='', style='->', color='#555555'):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle=style, color=color, lw=1.1))
        if label:
            ax.text((x1+x2)/2, (y1+y2)/2 + 0.13, label,
                    ha='center', fontsize=7, color=color,
                    bbox=dict(boxstyle='round,pad=0.1', fc='white', ec='none', alpha=0.85))

    return cls, enu, arr

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMME 7 : Diagramme de Classes Global — tous services
# ══════════════════════════════════════════════════════════════════════════════
def draw_classes():
    fig, ax = plt.subplots(figsize=(24, 18))
    ax.set_xlim(0, 24); ax.set_ylim(0, 18); ax.axis('off')
    fig.patch.set_facecolor('#f4f6f8')

    RH = 0.27; HH = 0.44

    def B(x, y, name, attrs, color, w=3.4):
        h = max(len(attrs)*RH+0.12, 0.32); tot = HH+h
        ax.add_patch(FancyBboxPatch((x,y-tot),w,tot,boxstyle='square',facecolor='white',edgecolor=color,lw=1.2))
        ax.add_patch(FancyBboxPatch((x,y-HH),w,HH,boxstyle='square',facecolor=color,edgecolor=color,lw=1.2))
        ax.text(x+w/2,y-HH/2,name,ha='center',va='center',fontsize=7.5,fontweight='bold',color='white')
        ax.plot([x,x+w],[y-HH,y-HH],color=color,lw=0.7)
        for i,a in enumerate(attrs):
            ax.text(x+0.09,y-HH-0.13-i*RH,f'  {a}',va='center',fontsize=6.0,color='#333333',family='monospace')

    def Z(x,y,w,h,label,color):
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle='round,pad=0.06',facecolor=(*color,0.06),edgecolor=color,lw=1.8,linestyle='--'))
        ax.text(x+w/2,y+h-0.16,f'<< {label} >>',ha='center',va='top',fontsize=7.5,fontweight='bold',color=color)

    CA=(0.50,0.20,0.60); CCL=(0.09,0.44,0.73); CAC=(0.85,0.40,0.10)
    CAG=(0.02,0.31,0.24); CLO=(0.06,0.62,0.49); CRE=(0.72,0.11,0.11)
    CTX=(0.00,0.48,0.68); CNO=(0.70,0.30,0.00); CRP=(0.38,0.38,0.38); CCO=(0.10,0.50,0.50)

    Z(0.1,12.2,11.8,5.5,'auth-service  :8080',CA)
    Z(12.1,12.2,5.8,5.5,'client-service  :8081',CCL)
    Z(18.1,12.2,5.8,5.5,'account-service  :8082',CAC)

    B(0.3,17.4,'User',['id:String(UUID)','email:String[unique]','userRoleType:UserRoleType','agencyId:String'],CA,w=3.3)
    B(3.8,17.4,'Role',['id:String','name:String[unique]','permissions:Set<Privilege>'],CA,w=2.9)
    B(6.9,17.4,'Privilege',['id:String','name:String[unique]','description:String'],CA,w=2.7)
    B(9.8,17.4,'RefreshToken',['token:String[unique]','user:User','revoked:boolean'],CA,w=1.9)
    B(0.3,14.9,'UserSession',['sessionToken:String','user:User','isActive:boolean'],CA,w=3.3)
    B(3.8,14.9,'AuditLog',['userId:String','action:String','status:String'],CA,w=3.2)
    B(7.2,14.9,'PasswordResetToken',['token:String[unique]','user:User','used:boolean'],CA,w=4.4)

    B(12.3,17.4,'Client',['id:String(UUID)','email:String','clientType:ClientType','status:ClientStatus','creditScore:Integer'],CCL,w=5.4)
    B(12.3,14.9,'Document',['type:DocumentType','fileUrl:String','verificationStatus:VerificationStatus'],CCL,w=5.4)

    B(18.3,17.4,'Compte',['id:Long','clientId:String','typeCompte:TypeCompte','solde:BigDecimal','statut:StatutCompte'],CAC,w=5.4)

    Z(0.1,7.1,5.8,4.9,'agency-service  :8086',CAG)
    Z(6.1,7.1,11.8,4.9,'loan-service  :8083',CLO)
    Z(18.1,7.1,5.8,4.9,'repayment-service  :8084',CRE)

    B(0.3,11.7,'Agency',['code:String[unique]','name:String','directorId:String'],CAG,w=2.6)
    B(3.1,11.7,'AgencyStats',['agencyId:String','totalClients:Long','totalLoans:Long'],CAG,w=2.6)
    B(0.3,10.1,'AgentAssignment',['agentId:String','agencyId:String','active:boolean'],CAG,w=2.6)
    B(3.1,10.1,'AgencyConfiguration',['agencyId:String','configKey:String','configValue:String'],CAG,w=2.6)

    B(6.3,11.7,'LoanProduct',['name:String[unique]','minAmount:BigDecimal','interestRate:BigDecimal'],CLO,w=3.7)
    B(10.2,11.7,'LoanApplication',['clientId:String','requestedAmount:BigDecimal','status:ApplicationStatus'],CLO,w=3.7)
    B(14.1,11.7,'Loan',['loanNumber:String','clientId:String','amount:BigDecimal','status:LoanStatus'],CLO,w=3.6)
    B(6.3,9.8,'AmortizationSchedule',['installmentNumber:Integer','dueAmount:BigDecimal','paid:boolean'],CLO,w=5.4)
    B(11.9,9.8,'Schedule(loan)',['loanId:String','dueDate:LocalDateTime','totalAmount:BigDecimal','status:ScheduleStatus'],CLO,w=5.8)

    B(18.3,11.7,'Repayment',['loanId:String','dueAmount:BigDecimal','status:PaymentStatus'],CRE,w=2.6)
    B(21.1,11.7,'Payment',['loanId:String','amount:BigDecimal','paymentMethod:PaymentMethod'],CRE,w=2.6)
    B(18.3,10.1,'Schedule(repay)',['loanId:String','dueAmount:BigDecimal','paid:boolean'],CRE,w=2.6)
    B(21.1,10.1,'Penalty',['loanId:String','amount:BigDecimal','daysOverdue:Integer'],CRE,w=2.6)

    Z(0.1,2.6,5.8,4.3,'transaction-service  :8088',CTX)
    Z(6.1,2.6,8.8,4.3,'notification-service  :8089',CNO)
    Z(15.1,2.6,8.8,4.3,'reporting-service  :8085',CRP)

    B(0.3,6.6,'Transaction',['compteId:Long','typeTransaction:TypeTransaction','montant:BigDecimal','statut:StatutTransaction','modePaiement:ModePaiement'],CTX,w=5.4)
    B(6.3,6.6,'Notification',['clientId:String','type:TypeNotification','canal:CanalNotification','statut:StatutNotification'],CNO,w=4.2)
    B(10.7,6.6,'NotificationTemplate',['nom:String[unique]','typeNotification:TypeNotification','canal:CanalNotification','actif:Boolean'],CNO,w=3.9)
    B(15.3,6.6,'Report',['name:String','type:ReportType','format:ReportFormat','generatedBy:String'],CRP,w=4.1)
    B(19.6,6.6,'Kpi',['name:String','category:String','value:BigDecimal','unit:String'],CRP,w=4.1)

    Z(0.1,0.1,23.8,2.3,'configuration-service  :8087',CCO)
    B(0.3,2.1,'MicrofinanceConfiguration',['microfinanceCode:String','clientIdStrategy:ClientIdGenerationStrategy'],CCO,w=6.2)
    B(6.7,2.1,'AccountCategory',['code:String[unique]','name:String'],CCO,w=5.0)
    B(12.0,2.1,'AccountTypeConfiguration',['code:String','accountType:AccountType','category:AccountCategory','interestRate:BigDecimal'],CCO,w=5.5)

    # relations intra-service
    ax.plot([3.6,3.8],[16.6,16.6],'-',color=CA,lw=1.1); ax.text(3.7,16.75,'*:*',fontsize=6,color=CA,ha='center')
    ax.plot([6.7,6.9],[16.6,16.6],'-',color=CA,lw=1.1); ax.text(6.8,16.75,'*:*',fontsize=6,color=CA,ha='center')
    ax.annotate('',xy=(15.0,14.9),xytext=(15.0,16.0),arrowprops=dict(arrowstyle='->',color=CCL,lw=1.0))
    ax.text(15.3,15.5,'1..*',fontsize=6,color=CCL)
    ax.annotate('',xy=(9.0,9.8),xytext=(15.9,10.8),arrowprops=dict(arrowstyle='->',color=CLO,lw=1.0))
    ax.text(12.2,10.4,'1..*',fontsize=6,color=CLO,ha='center')
    ax.plot([11.7,12.0],[1.4,1.4],'-',color=CCO,lw=1.1); ax.text(11.85,1.55,'1..*',fontsize=6,color=CCO,ha='center')
    ax.plot([10.5,10.7],[5.5,5.5],'-',color=CNO,lw=1.1); ax.text(10.6,5.65,'N:1',fontsize=6,color=CNO,ha='center')

    note=('Relations cross-services (Database per Service) :  '
          'Compte.clientId→Client.id  |  Transaction.compteId→Compte.id  |  '
          'Loan.clientId→Client.id  |  Repayment.loanId→Loan.id')
    ax.text(12.0,0.04,note,ha='center',va='bottom',fontsize=6.0,color='#555555',style='italic',
            bbox=dict(boxstyle='round,pad=0.2',fc='#fffbe6',ec='#ccaa00',lw=0.8))

    ax.set_title('Diagramme de Classes Global — MicrofinanceHub  (10 microservices)',
                 fontsize=14, fontweight='bold', color='#064e3b', pad=12)
    return fig_to_stream(fig)

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMMES 8a-8j : Diagrammes de Classes par service
# ══════════════════════════════════════════════════════════════════════════════

def draw_classes_auth_service():
    fig, ax = plt.subplots(figsize=(20, 14))
    ax.set_xlim(0, 20); ax.set_ylim(0, 14); ax.axis('off')
    fig.patch.set_facecolor('#f4f6f8')
    cls, enu, arr = _uml(ax)
    C = (0.50, 0.20, 0.60)
    cls(0.2, 13.5, 'User', [
        '- id : String (UUID)', '- email : String [unique]',
        '- phoneNumber : String', '- password : String (BCrypt)',
        '- firstName : String', '- lastName : String',
        '- userRoleType : UserRoleType', '- agencyId : String',
        '- agencyCode : String', '- enabled : boolean',
        '- locked : boolean', '- createdAt : LocalDateTime',
        '- lastLoginAt : LocalDateTime'], C, w=4.8)
    cls(5.3, 13.5, 'Role', [
        '- id : String (UUID)', '- name : String [unique]',
        '- description : String', '- permissions : Set<Privilege>'], C, w=4.0)
    cls(9.6, 13.5, 'Privilege', [
        '- id : String (UUID)', '- name : String [unique]',
        '- description : String'], C, w=4.0)
    cls(14.0, 13.5, 'RefreshToken', [
        '- id : String (UUID)', '- token : String [unique]',
        '- user : User', '- expiryDate : LocalDateTime',
        '- revoked : boolean'], C, w=5.7)
    cls(0.2, 8.2, 'UserSession', [
        '- id : String (UUID)', '- user : User',
        '- sessionToken : String [unique]', '- deviceInfo : String',
        '- ipAddress : String', '- lastActivity : LocalDateTime',
        '- expiresAt : LocalDateTime', '- isActive : boolean'], C, w=4.8)
    cls(5.3, 8.2, 'AuditLog', [
        '- id : String (UUID)', '- userId : String',
        '- email : String', '- action : String',
        '- ipAddress : String', '- userAgent : String',
        '- status : String', '- failureReason : String',
        '- timestamp : LocalDateTime'], C, w=4.8)
    cls(10.4, 8.2, 'PasswordResetToken', [
        '- id : String (UUID)', '- user : User',
        '- token : String [unique]', '- expiryDate : LocalDateTime',
        '- used : boolean'], C, w=5.0)
    enu(0.2, 3.8, 'UserRoleType', [
        'CLIENT', 'AGENT', 'DIRECTEUR_AGENCE', 'ADMIN'], C, w=4.0)
    enu(4.6, 3.8, 'Permission', [
        'VIEW_OWN_PROFILE', 'UPDATE_OWN_PROFILE', 'APPLY_FOR_LOAN',
        'VIEW_OWN_LOANS', 'MAKE_REPAYMENT', 'VIEW_OWN_ACCOUNTS',
        'VIEW_ALL_CLIENTS', 'MANAGE_CLIENTS', 'APPROVE_LOANS',
        'VIEW_BASIC_REPORTS', 'MANAGE_USERS', 'MANAGE_ROLES',
        'CONFIGURE_SYSTEM', 'VIEW_ALL_REPORTS', 'VIEW_AUDIT_LOGS'], C, w=5.5)
    arr(5.0, 12.4, 5.3, 12.4, 'User *:* Role', color=C)
    arr(9.3, 12.4, 9.6, 12.4, 'Role *:* Privilege', color=C)
    arr(14.3, 12.0, 5.0, 11.8, '1:1', color=C)
    arr(2.6, 8.2, 2.6, 8.9, 'N:1', color=C)
    arr(12.9, 8.2, 12.9, 8.9, '1:1', color=C)
    ax.set_title('Diagramme de Classes — auth-service (:8080)',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)


def draw_classes_client_service():
    fig, ax = plt.subplots(figsize=(14, 10))
    ax.set_xlim(0, 14); ax.set_ylim(0, 10); ax.axis('off')
    fig.patch.set_facecolor('#f4f6f8')
    cls, enu, arr = _uml(ax)
    C = (0.09, 0.44, 0.73)
    cls(0.2, 9.5, 'Client', [
        '- id : String (UUID)', '- email : String [unique]',
        '- phoneNumber : String', '- firstName : String',
        '- lastName : String', '- address : String',
        '- birthDate : LocalDateTime', '- clientType : ClientType',
        '- status : ClientStatus', '- creditScore : Integer',
        '- agencyId : String', '- createdBy : String',
        '- createdAt : LocalDateTime'], C, w=6.0)
    cls(6.5, 9.5, 'Document', [
        '- id : String (UUID)', '- client : Client',
        '- type : DocumentType', '- fileName : String',
        '- fileUrl : String', '- fileType : String',
        '- fileSize : Long', '- verificationStatus : VerificationStatus',
        '- verifiedBy : String', '- verifiedAt : LocalDateTime',
        '- uploadedAt : LocalDateTime'], C, w=7.2)
    enu(0.2, 4.0, 'ClientType', ['INDIVIDUAL', 'BUSINESS', 'GROUP'], C, w=3.2)
    enu(3.6, 4.0, 'ClientStatus', [
        'ACTIVE', 'INACTIVE', 'SUSPENDED', 'BLACKLISTED', 'PENDING'], C, w=3.2)
    enu(7.0, 4.0, 'DocumentType', [
        'ID_CARD', 'PASSPORT', 'DRIVER_LICENSE',
        'PROOF_OF_ADDRESS', 'BUSINESS_REGISTRATION',
        'TAX_IDENTIFICATION', 'BANK_STATEMENT', 'PHOTO'], C, w=3.6)
    enu(10.8, 4.0, 'VerificationStatus', [
        'PENDING', 'VERIFIED', 'REJECTED', 'EXPIRED'], C, w=3.0)
    arr(6.2, 7.0, 6.5, 7.0, '1 — *', color=C)
    ax.set_title('Diagramme de Classes — client-service (:8081)',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)


def draw_classes_account_service():
    fig, ax = plt.subplots(figsize=(12, 10))
    ax.set_xlim(0, 12); ax.set_ylim(0, 10); ax.axis('off')
    fig.patch.set_facecolor('#f4f6f8')
    cls, enu, arr = _uml(ax)
    C = (0.85, 0.40, 0.10)
    cls(0.5, 9.5, 'Compte', [
        '- id : Long', '- clientId : String',
        '- numeroCompte : String [unique]', '- typeCompte : TypeCompte',
        '- solde : BigDecimal', '- devise : Devise',
        '- statut : StatutCompte', '- dateOuverture : LocalDateTime',
        '- tauxInteret : BigDecimal', '- soldeMinimum : BigDecimal',
        '- plafond : BigDecimal', '- description : String',
        '- clientEmail : String', '- clientNom : String',
        '+ crediter(montant)', '+ debiter(montant)', '+ isOperationnel()'], C, w=11.0)
    enu(0.3, 3.8, 'TypeCompte', [
        'EPARGNE', 'COURANT', 'DEPOT_A_TERME', 'MICRO_EPARGNE', 'CREDIT'], C, w=3.5)
    enu(4.1, 3.8, 'StatutCompte', [
        'EN_ATTENTE_VALIDATION', 'ACTIF', 'BLOQUE',
        'SUSPENDU', 'INACTIF', 'FERME', 'REJETE'], C, w=4.0)
    enu(8.3, 3.8, 'Devise', [
        'XAF', 'XOF', 'EUR', 'USD', 'GBP', 'NGN', 'GHS', 'MAD', 'CDF'], C, w=3.5)
    ax.set_title('Diagramme de Classes — account-service (:8082)',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)


def draw_classes_agency_service():
    fig, ax = plt.subplots(figsize=(16, 10))
    ax.set_xlim(0, 16); ax.set_ylim(0, 10); ax.axis('off')
    fig.patch.set_facecolor('#f4f6f8')
    cls, enu, arr = _uml(ax)
    C = (0.02, 0.31, 0.24)
    cls(0.2, 9.5, 'Agency', [
        '- id : String (UUID)', '- code : String [unique]',
        '- name : String', '- address : String', '- city : String',
        '- phoneNumber : String', '- email : String',
        '- directorId : String [unique]', '- directorEmail : String',
        '- directorName : String', '- region : String',
        '- status : String', '- createdAt : LocalDateTime'], C, w=5.5)
    cls(6.0, 9.5, 'AgentAssignment', [
        '- id : String (UUID)', '- agentId : String',
        '- agentEmail : String', '- agentName : String',
        '- agencyId : String', '- agencyCode : String',
        '- role : String', '- assignedBy : String',
        '- reason : String', '- reference : String [unique]',
        '- assignedAt : LocalDateTime', '- active : boolean'], C, w=5.5)
    cls(11.8, 9.5, 'AgencyStats', [
        '- agencyId : String (PK)', '- totalClients : Long',
        '- totalAccounts : Long', '- totalLoans : Long',
        '- activeLoans : Long', '- completedLoans : Long',
        '- defaultedLoans : Long', '- totalOutstanding : BigDecimal',
        '- monthlyRepaymentRate : Double'], C, w=3.9)
    cls(0.2, 3.5, 'AgencyConfiguration', [
        '- id : String (UUID)', '- agencyId : String',
        '- configKey : String', '- configValue : String',
        '- description : String', '- updatedBy : String',
        '- createdAt : LocalDateTime'], C, w=7.0)
    arr(5.5, 8.0, 6.0, 8.0, '1 — *', color=C)
    arr(5.5, 7.2, 11.8, 8.0, '1 — 1', color=C)
    arr(3.7, 5.8, 3.7, 6.5, '1 — *', color=C)
    ax.set_title('Diagramme de Classes — agency-service (:8086)',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)


def draw_classes_loan_service():
    fig, ax = plt.subplots(figsize=(22, 14))
    ax.set_xlim(0, 22); ax.set_ylim(0, 14); ax.axis('off')
    fig.patch.set_facecolor('#f4f6f8')
    cls, enu, arr = _uml(ax)
    C = (0.06, 0.62, 0.49)
    cls(0.2, 13.5, 'LoanProduct', [
        '- id : String (UUID)', '- name : String [unique]',
        '- description : String', '- minAmount : BigDecimal',
        '- maxAmount : BigDecimal', '- minTermMonths : Integer',
        '- maxTermMonths : Integer', '- interestRate : BigDecimal',
        '- active : boolean'], C, w=5.0)
    cls(5.5, 13.5, 'LoanApplication', [
        '- id : String (UUID)', '- applicationNumber : String',
        '- clientId : String', '- accountNumber : String',
        '- clientEmail : String', '- requestedAmount : BigDecimal',
        '- termMonths : Integer', '- purpose : String',
        '- monthlyIncome : BigDecimal', '- employmentStatus : String',
        '- status : ApplicationStatus', '- rejectionReason : String',
        '- reviewedBy : String', '- applicationDate : LocalDateTime'], C, w=5.5)
    cls(11.3, 13.5, 'Loan', [
        '- id : String (UUID)', '- loanNumber : String',
        '- applicationId : String', '- clientId : String',
        '- amount : BigDecimal', '- interestRate : BigDecimal',
        '- termMonths : Integer', '- repaymentFrequency : RepaymentFrequency',
        '- monthlyPayment : BigDecimal', '- totalRepayment : BigDecimal',
        '- remainingBalance : BigDecimal', '- status : LoanStatus',
        '- loanProductId : String', '- approvedBy : String',
        '- disbursementDate : LocalDateTime'], C, w=5.5)
    cls(17.0, 13.5, 'AmortizationSchedule', [
        '- id : String (UUID)', '- loan : Loan',
        '- installmentNumber : Integer', '- dueDate : LocalDateTime',
        '- dueAmount : BigDecimal', '- principalAmount : BigDecimal',
        '- interestAmount : BigDecimal', '- remainingBalance : BigDecimal',
        '- paid : boolean', '- paidDate : LocalDateTime',
        '- paymentId : String'], C, w=4.8)
    cls(0.2, 5.5, 'Schedule', [
        '- id : String (UUID)', '- loanId : String',
        '- loanNumber : String', '- clientId : String',
        '- installmentNumber : Integer', '- dueDate : LocalDateTime',
        '- principalAmount : BigDecimal', '- interestAmount : BigDecimal',
        '- totalAmount : BigDecimal', '- paidAmount : BigDecimal',
        '- remainingAmount : BigDecimal', '- penaltyAmount : BigDecimal',
        '- status : ScheduleStatus', '- paidDate : LocalDateTime',
        '- paymentId : String'], C, w=5.5)
    enu(6.5, 5.0, 'LoanStatus', [
        'PENDING', 'UNDER_REVIEW', 'APPROVED',
        'REJECTED', 'ACTIVE', 'COMPLETED', 'DEFAULTED'], C, w=4.0)
    enu(10.8, 5.0, 'ApplicationStatus', [
        'PENDING', 'UNDER_REVIEW', 'APPROVED', 'REJECTED', 'CANCELLED'], C, w=4.0)
    enu(15.1, 5.0, 'RepaymentFrequency', [
        'WEEKLY', 'BIWEEKLY', 'MONTHLY', 'QUARTERLY'], C, w=3.5)
    enu(18.9, 5.0, 'ScheduleStatus', [
        'PENDING', 'PARTIALLY_PAID', 'PAID', 'OVERDUE'], C, w=2.9)
    arr(11.0, 11.5, 11.3, 11.5, '1 — *', color=C)
    arr(16.8, 10.5, 17.0, 10.5, '1 — *', color=C)
    arr(13.5, 10.5, 3.5, 5.5, '1 — *', color=C)
    ax.set_title('Diagramme de Classes — loan-service (:8083)',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)


def draw_classes_repayment_service():
    fig, ax = plt.subplots(figsize=(18, 12))
    ax.set_xlim(0, 18); ax.set_ylim(0, 12); ax.axis('off')
    fig.patch.set_facecolor('#f4f6f8')
    cls, enu, arr = _uml(ax)
    C = (0.72, 0.11, 0.11)
    cls(0.2, 11.5, 'Repayment', [
        '- id : String (UUID)', '- loanId : String',
        '- clientId : String', '- installmentNumber : Integer',
        '- dueAmount : BigDecimal', '- paidAmount : BigDecimal',
        '- penaltyAmount : BigDecimal', '- dueDate : LocalDateTime',
        '- paidDate : LocalDateTime', '- status : PaymentStatus',
        '- paymentId : String'], C, w=5.5)
    cls(6.0, 11.5, 'Payment', [
        '- id : String (UUID)', '- paymentNumber : String',
        '- loanId : String', '- clientId : String',
        '- amount : BigDecimal', '- penaltyAmount : BigDecimal',
        '- paymentMethod : PaymentMethod', '- status : PaymentStatus',
        '- transactionId : String', '- reference : String',
        '- receiptNumber : String', '- paymentDate : LocalDateTime',
        '- compteSourceId : Long', '- validatedBy : String'], C, w=5.5)
    cls(11.8, 11.5, 'Schedule', [
        '- id : UUID', '- loanId : String',
        '- installmentNumber : Integer', '- dueDate : LocalDateTime',
        '- dueAmount : BigDecimal', '- principalAmount : BigDecimal',
        '- interestAmount : BigDecimal', '- remainingBalance : BigDecimal',
        '- paid : boolean', '- paidDate : LocalDateTime',
        '- paymentId : String',
        '+ getPaidAmount()', '+ getRemainingAmount()', '+ getStatus()'], C, w=5.9)
    cls(0.2, 4.5, 'Penalty', [
        '- id : String (UUID)', '- loanId : String',
        '- scheduleId : String', '- installmentNumber : Integer',
        '- amount : BigDecimal', '- dailyPenaltyRate : BigDecimal',
        '- daysOverdue : Integer', '- paid : boolean',
        '- appliedDate : LocalDateTime', '- paidDate : LocalDateTime',
        '- paymentId : String'], C, w=5.5)
    enu(6.5, 4.0, 'PaymentMethod', [
        'ESPECES', 'MOBILE_MONEY', 'VIREMENT_BANCAIRE',
        'VIREMENT_INTERNE', 'CHEQUE', 'DEBIT_COMPTE'], C, w=4.0)
    enu(10.8, 4.0, 'PaymentStatus', [
        'PENDING', 'PENDING_VALIDATION',
        'COMPLETED', 'FAILED', 'CANCELLED'], C, w=4.0)
    enu(15.1, 4.0, 'ScheduleStatus', [
        'PENDING', 'PARTIALLY_PAID', 'PAID', 'OVERDUE'], C, w=2.7)
    ax.set_title('Diagramme de Classes — repayment-service (:8084)',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)


def draw_classes_transaction_service():
    fig, ax = plt.subplots(figsize=(14, 10))
    ax.set_xlim(0, 14); ax.set_ylim(0, 10); ax.axis('off')
    fig.patch.set_facecolor('#f4f6f8')
    cls, enu, arr = _uml(ax)
    C = (0.00, 0.48, 0.68)
    cls(0.5, 9.5, 'Transaction', [
        '- id : Long', '- compteId : Long',
        '- typeTransaction : TypeTransaction', '- montant : BigDecimal',
        '- soldeAvant : BigDecimal', '- soldeApres : BigDecimal',
        '- dateTransaction : LocalDateTime', '- reference : String [unique]',
        '- statut : StatutTransaction', '- description : String',
        '- agentId : Long', '- compteDestination : String',
        '- loanId : String', '- clientId : String',
        '- modePaiement : ModePaiement', '- pieceJustificative : String',
        '- campayReference : String [unique]', '- numeroPaiement : String',
        '- createdAt : LocalDateTime'], C, w=13.0)
    enu(0.3, 3.8, 'TypeTransaction', [
        'DEPOT', 'RETRAIT', 'VIREMENT_ENTRANT', 'VIREMENT_SORTANT',
        'DECAISSEMENT_PRET', 'REMBOURSEMENT_PRET', 'FRAIS', 'INTERET'], C, w=4.2)
    enu(4.8, 3.8, 'StatutTransaction', [
        'INITIEE', 'EN_VALIDATION', 'VALIDEE',
        'EN_TRAITEMENT', 'COMPLETEE', 'ECHOUEE', 'ANNULEE'], C, w=4.2)
    enu(9.3, 3.8, 'ModePaiement', [
        'ESPECES', 'MOBILE_MONEY', 'VIREMENT_BANCAIRE',
        'VIREMENT_INTERNE', 'CHEQUE', 'DEBIT_COMPTE'], C, w=4.5)
    ax.set_title('Diagramme de Classes — transaction-service (:8088)',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)


def draw_classes_notification_service():
    fig, ax = plt.subplots(figsize=(16, 12))
    ax.set_xlim(0, 16); ax.set_ylim(0, 12); ax.axis('off')
    fig.patch.set_facecolor('#f4f6f8')
    cls, enu, arr = _uml(ax)
    C = (0.70, 0.30, 0.00)
    cls(0.2, 11.5, 'Notification', [
        '- id : Long', '- clientId : String',
        '- destinataireEmail : String', '- destinataireTelephone : String',
        '- sujet : String', '- message : String',
        '- type : TypeNotification', '- canal : CanalNotification',
        '- statut : StatutNotification', '- priorite : Integer',
        '- dateEnvoi : LocalDateTime', '- dateProgrammee : LocalDateTime',
        '- tentatives : Integer', '- erreurMessage : String',
        '- referenceId : String', '- referenceType : String',
        '- template : NotificationTemplate'], C, w=7.5)
    cls(8.0, 11.5, 'NotificationTemplate', [
        '- id : Long', '- nom : String [unique]',
        '- sujet : String', '- contenu : String (variables {{}})',
        '- typeNotification : TypeNotification',
        '- canal : CanalNotification',
        '- actif : Boolean', '- fichierHtml : String'], C, w=7.7)
    enu(0.2, 4.5, 'TypeNotification', [
        'DEMANDE_PRET', 'APPROBATION_PRET', 'REJET_PRET', 'DECAISSEMENT',
        'RAPPEL_ECHEANCE', 'CONFIRMATION_REMB', 'ALERTE_RETARD', 'PENALITE_APPLIQUEE',
        'CREATION_COMPTE', 'DEPOT_EFFECTUE', 'RETRAIT_EFFECTUE',
        'PROMOTION', 'NEWSLETTER', 'ALERTE_SYSTEME', 'ALERTE_SOLDE'], C, w=5.0)
    enu(5.5, 4.5, 'CanalNotification', [
        'EMAIL', 'SMS', 'IN_APP', 'EMAIL_SMS'], C, w=3.5)
    enu(9.3, 4.5, 'StatutNotification', [
        'EN_ATTENTE', 'EN_COURS', 'ENVOYEE',
        'ECHEC', 'ECHEC_DEFINITIF', 'PROGRAMMEE', 'LUE'], C, w=4.0)
    arr(7.7, 9.0, 8.0, 9.0, 'N — 1', color=C)
    ax.set_title('Diagramme de Classes — notification-service (:8089)',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)


def draw_classes_reporting_service():
    fig, ax = plt.subplots(figsize=(14, 10))
    ax.set_xlim(0, 14); ax.set_ylim(0, 10); ax.axis('off')
    fig.patch.set_facecolor('#f4f6f8')
    cls, enu, arr = _uml(ax)
    C = (0.38, 0.38, 0.38)
    cls(0.2, 9.5, 'Report', [
        '- id : String (UUID)', '- name : String',
        '- description : String', '- type : ReportType',
        '- format : ReportFormat', '- startDate : LocalDateTime',
        '- endDate : LocalDateTime', '- filePath : String',
        '- fileSize : Long', '- generatedBy : String',
        '- generatedAt : LocalDateTime', '- scheduled : boolean',
        '- scheduleCron : String'], C, w=6.0)
    cls(6.5, 9.5, 'Kpi', [
        '- id : String (UUID)', '- name : String',
        '- description : String', '- category : String',
        '- value : BigDecimal', '- unit : String',
        '- periodStart : LocalDateTime', '- periodEnd : LocalDateTime',
        '- calculatedAt : LocalDateTime', '- calculatedBy : String',
        '- metadata : String (jsonb)'], C, w=7.2)
    enu(0.2, 3.5, 'ReportType', [
        'PORTFOLIO_SUMMARY', 'LOAN_PERFORMANCE',
        'CLIENT_ANALYSIS', 'REPAYMENT_HISTORY',
        'AGENCY_PERFORMANCE', 'FINANCIAL_STATEMENT'], C, w=4.5)
    enu(5.0, 3.5, 'ReportFormat', ['PDF', 'EXCEL', 'CSV', 'JSON'], C, w=3.0)
    enu(8.3, 3.5, 'PeriodType', [
        'DAILY', 'WEEKLY', 'MONTHLY', 'QUARTERLY', 'YEARLY', 'CUSTOM'], C, w=3.5)
    ax.set_title('Diagramme de Classes — reporting-service (:8085)',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)


def draw_classes_configuration_service():
    fig, ax = plt.subplots(figsize=(16, 12))
    ax.set_xlim(0, 16); ax.set_ylim(0, 12); ax.axis('off')
    fig.patch.set_facecolor('#f4f6f8')
    cls, enu, arr = _uml(ax)
    C = (0.10, 0.50, 0.50)
    cls(0.2, 11.5, 'MicrofinanceConfiguration', [
        '- id : String (UUID)', '- microfinanceCode : String [unique]',
        '- affiliatedBankCode : String', '- affiliatedToBank : boolean',
        '- agencyCode : String', '- agencyName : String',
        '- clientIdStrategy : ClientIdGenerationStrategy',
        '- customClientIdPattern : String', '- enableRibGeneration : boolean',
        '- bankCode : String', '- countryCode : String',
        '- accountNumberFormat : String', '- useCustomFormat : boolean',
        '- active : boolean', '- createdBy : String'], C, w=7.5)
    cls(8.0, 11.5, 'AccountCategory', [
        '- id : String (UUID)', '- code : String [unique]',
        '- name : String', '- description : String',
        '- icon : String', '- color : String',
        '- displayOrder : int', '- active : boolean',
        '- accountTypes : List<AccountTypeConfiguration>',
        '- createdBy : String'], C, w=7.7)
    cls(0.2, 4.5, 'AccountTypeConfiguration', [
        '- id : String (UUID)', '- code : String [unique]',
        '- accountType : AccountType', '- category : AccountCategory',
        '- name : String', '- description : String',
        '- minimumBalance : BigDecimal', '- maximumBalance : BigDecimal',
        '- interestRate : BigDecimal', '- monthlyFee : BigDecimal',
        '- transactionFee : BigDecimal', '- allowOverdraft : boolean',
        '- overdraftLimit : BigDecimal', '- maxAccountsPerClient : int',
        '- active : boolean'], C, w=7.5)
    enu(8.2, 4.5, 'AccountType', [
        'EPARGNE', 'COURANT', 'DEPOT_A_TERME', 'MICRO_EPARGNE', 'CREDIT'], C, w=4.0)
    enu(12.5, 4.5, 'ClientIdGenerationStrategy', [
        'SEQUENTIAL', 'UUID', 'CUSTOM'], C, w=3.3)
    ax.plot([7.9, 7.9, 7.7], [9.2, 3.5, 3.5], '-', color=C, lw=1.2)
    ax.text(7.5, 6.3, '1', fontsize=7, color=C); ax.text(7.5, 3.3, '*', fontsize=9, color=C)
    ax.set_title('Diagramme de Classes — configuration-service (:8087)',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)


def draw_classes_par_service():
    return draw_classes_auth_service()

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMME 9 : États/transitions d'un compte bancaire — NOUVEAU
# ══════════════════════════════════════════════════════════════════════════════
def draw_state_compte():
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14); ax.set_ylim(0, 8); ax.axis('off')
    fig.patch.set_facecolor('#f8f9fa')

    # Couleurs des états
    col_attente = (0.7, 0.55, 0.0)
    col_actif   = VERT
    col_bloque  = ORANGE
    col_ferme   = (0.4, 0.4, 0.4)
    col_rejete  = ROUGE

    def state(x, y, label, col, w=2.2, h=0.75):
        ax.add_patch(FancyBboxPatch((x - w/2, y - h/2), w, h,
                     boxstyle="round,pad=0.12",
                     facecolor=(*col, 0.85), edgecolor=col, lw=2))
        ax.text(x, y, label, ha='center', va='center',
                fontsize=9, fontweight='bold', color='white')

    def transition(x1, y1, x2, y2, label, col='#444444',
                   conn='arc3,rad=0.0', lx=None, ly=None):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle='->', color=col, lw=1.5,
                            connectionstyle=conn))
        mlx = (x1+x2)/2 if lx is None else lx
        mly = (y1+y2)/2 + 0.18 if ly is None else ly
        ax.text(mlx, mly, label, ha='center', va='bottom',
                fontsize=7.5, color=col,
                bbox=dict(boxstyle='round,pad=0.15', facecolor='white',
                          edgecolor=col, lw=0.7, alpha=0.85))

    # État initial
    ax.add_patch(plt.Circle((2.0, 7.0), 0.22, color='black', zorder=5))

    # États
    state(5.0, 7.0, 'EN_ATTENTE\nVALIDATION', col_attente)
    state(9.0, 7.0, 'ACTIF', col_actif)
    state(11.5, 4.5, 'BLOQUE', col_bloque)
    state(9.0, 2.0, 'FERME', col_ferme)
    state(5.0, 2.0, 'REJETE', col_rejete)

    # État final (double cercle autour de FERME)
    ax.add_patch(plt.Circle((9.0, 2.0), 0.62, color=col_ferme, fill=False, lw=2.5))

    # Transitions
    transition(2.22, 7.0, 3.88, 7.0, 'creation()')
    transition(6.12, 7.0, 7.88, 7.0, 'approbation()\n[directeur/admin]', col=col_actif)
    transition(9.0, 6.62, 9.0, 6.62, '', col='white')  # placeholder
    transition(10.12, 7.0, 10.72, 5.25, 'blocage()\n[admin/directeur]',
               col=ORANGE, conn='arc3,rad=-0.2', lx=11.2, ly=6.4)
    transition(10.88, 4.12, 9.62, 6.62, 'deblocage()\n[admin]',
               col=VERT, conn='arc3,rad=-0.2', lx=8.8, ly=5.5)
    transition(9.0, 6.62, 9.0, 2.38, 'fermeture()\n[admin]',
               col=col_ferme, conn='arc3,rad=0.3', lx=10.5, ly=4.5)
    transition(3.88, 7.0, 5.95, 2.38, 'rejet()\n[admin/directeur]',
               col=col_rejete, conn='arc3,rad=0.3', lx=3.8, ly=4.5)

    # Légende
    ax.text(7.0, 0.4, 'Etat initial : cercle noir  |  Etat final : double cercle  |  Transitions : fleches etiquetees',
            ha='center', fontsize=7.5, style='italic', color='#666666')

    ax.set_title('Diagramme d\'Etats — Compte Bancaire (StatutCompte)',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMME 10 : États/transitions d'un prêt — NOUVEAU
# ══════════════════════════════════════════════════════════════════════════════
def draw_state_loan():
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14); ax.set_ylim(0, 8); ax.axis('off')
    fig.patch.set_facecolor('#f8f9fa')

    col_pending   = (0.7, 0.55, 0.0)
    col_approved  = BLEU
    col_rejected  = ROUGE
    col_active    = VERT
    col_completed = (0.1, 0.5, 0.2)
    col_defaulted = (0.6, 0.1, 0.5)

    def state(x, y, label, col, w=2.4, h=0.78):
        ax.add_patch(FancyBboxPatch((x - w/2, y - h/2), w, h,
                     boxstyle="round,pad=0.12",
                     facecolor=(*col, 0.85), edgecolor=col, lw=2))
        ax.text(x, y, label, ha='center', va='center',
                fontsize=9, fontweight='bold', color='white')

    def tr(x1, y1, x2, y2, label, col='#444444',
           conn='arc3,rad=0.0', lx=None, ly=None):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle='->', color=col, lw=1.5,
                            connectionstyle=conn))
        mlx = (x1+x2)/2 if lx is None else lx
        mly = (y1+y2)/2 + 0.2 if ly is None else ly
        ax.text(mlx, mly, label, ha='center', va='bottom',
                fontsize=7.5, color=col,
                bbox=dict(boxstyle='round,pad=0.15', facecolor='white',
                          edgecolor=col, lw=0.7, alpha=0.85))

    # État initial
    ax.add_patch(plt.Circle((1.2, 6.0), 0.22, color='black', zorder=5))

    # États
    state(4.0, 6.0,  'PENDING',   col_pending)
    state(7.5, 7.2,  'APPROVED',  col_approved)
    state(7.5, 4.8,  'REJECTED',  col_rejected)
    state(11.0, 7.2, 'ACTIVE',    col_active)
    state(11.0, 4.8, 'COMPLETED', col_completed)
    state(11.0, 2.5, 'DEFAULTED', col_defaulted)

    # État final (double cercle)
    ax.add_patch(plt.Circle((11.0, 4.8), 0.62, color=col_completed, fill=False, lw=2.5))

    # Transitions
    tr(1.42, 6.0, 2.78, 6.0, 'demande()')
    tr(5.2, 6.3, 6.32, 7.0, 'approbation()\n[admin/directeur]', col=col_approved)
    tr(5.2, 5.7, 6.32, 5.0, 'rejet()\n[admin/directeur]', col=col_rejected)
    tr(8.72, 7.2, 9.78, 7.2, 'decaissement()\n[admin]', col=col_active)
    tr(11.0, 6.82, 11.0, 5.18, 'remboursement\ncomplet()', col=col_completed)
    tr(11.0, 4.42, 11.0, 2.89, 'defaut\npaiement()', col=col_defaulted)
    tr(9.78, 7.0, 5.2, 6.3, 'annulation()\n[admin]',
       col='#888888', conn='arc3,rad=0.3', lx=7.0, ly=7.8)

    ax.text(7.0, 0.4, 'Etat initial : cercle noir  |  Etat final : double cercle (COMPLETED)',
            ha='center', fontsize=7.5, style='italic', color='#666666')

    ax.set_title('Diagramme d\'Etats — Pret (LoanStatus)',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMME 11 : États/transitions d'un paiement — NOUVEAU
# ══════════════════════════════════════════════════════════════════════════════
def draw_state_payment():
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14); ax.set_ylim(0, 8); ax.axis('off')
    fig.patch.set_facecolor('#f8f9fa')

    col_pending     = (0.7, 0.55, 0.0)
    col_pend_val    = BLEU
    col_completed   = VERT
    col_failed      = ROUGE

    def state(x, y, label, col, w=2.6, h=0.82):
        ax.add_patch(FancyBboxPatch((x - w/2, y - h/2), w, h,
                     boxstyle="round,pad=0.12",
                     facecolor=(*col, 0.85), edgecolor=col, lw=2))
        ax.text(x, y, label, ha='center', va='center',
                fontsize=9, fontweight='bold', color='white')

    def tr(x1, y1, x2, y2, label, col='#444444',
           conn='arc3,rad=0.0', lx=None, ly=None):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle='->', color=col, lw=1.5,
                            connectionstyle=conn))
        mlx = (x1+x2)/2 if lx is None else lx
        mly = (y1+y2)/2 + 0.2 if ly is None else ly
        ax.text(mlx, mly, label, ha='center', va='bottom',
                fontsize=7.5, color=col,
                bbox=dict(boxstyle='round,pad=0.15', facecolor='white',
                          edgecolor=col, lw=0.7, alpha=0.85))

    # État initial
    ax.add_patch(plt.Circle((1.2, 5.5), 0.22, color='black', zorder=5))

    # États
    state(4.5, 6.5,  'PENDING',             col_pending)
    state(4.5, 4.5,  'PENDING\nVALIDATION', col_pend_val)
    state(9.5, 5.5,  'COMPLETED',           col_completed)
    state(9.5, 3.0,  'FAILED',              col_failed)

    # État final (double cercle autour de COMPLETED)
    ax.add_patch(plt.Circle((9.5, 5.5), 0.65, color=col_completed, fill=False, lw=2.5))

    # Transitions
    tr(1.42, 5.5, 3.18, 6.3, 'agent/directeur\nenregistre()', col=col_pending)
    tr(1.42, 5.5, 3.18, 4.7, 'client\nenregistre()', col=col_pend_val)

    tr(5.82, 6.5, 8.18, 5.7, 'admin valide()\n[PENDING -> COMPLETED]',
       col=col_completed, lx=7.0, ly=7.0)
    tr(5.82, 4.5, 8.18, 5.3, 'admin valide()\n[PENDING_VAL -> COMPLETED]',
       col=col_completed, lx=7.0, ly=4.6)

    tr(5.82, 6.3, 8.18, 3.3, 'echec()\n[timeout/erreur]',
       col=col_failed, conn='arc3,rad=0.2', lx=5.5, ly=4.5)
    tr(5.82, 4.3, 8.18, 3.1, 'echec()\n[rejet admin]',
       col=col_failed, conn='arc3,rad=-0.2', lx=8.0, ly=3.3)

    # Flux CamPay (webhook direct -> COMPLETED)
    ax.annotate('', xy=(8.18, 5.5), xytext=(1.42, 5.5),
        arrowprops=dict(arrowstyle='->', color='#888888', lw=1.3, linestyle='dashed',
                        connectionstyle='arc3,rad=-0.5'))
    ax.text(4.8, 2.8, 'CamPay webhook\n(direct -> COMPLETED)', ha='center',
            fontsize=7, color='#888888', style='italic',
            bbox=dict(boxstyle='round,pad=0.2', facecolor='#f5f5f5',
                      edgecolor='#aaaaaa', lw=0.8))

    ax.text(7.0, 0.4,
            'Etat initial : cercle noir  |  Etat final : COMPLETED (double cercle)  '
            '|  Flux CamPay = tirete',
            ha='center', fontsize=7, style='italic', color='#666666')

    ax.set_title('Diagramme d\'Etats — Paiement/Remboursement (PaymentStatus)',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMME 12 : Communication inter-services (RabbitMQ)
# ══════════════════════════════════════════════════════════════════════════════
def draw_rabbitmq():
    fig, ax = plt.subplots(figsize=(14, 8))
    ax.set_xlim(0, 14); ax.set_ylim(0, 8); ax.axis('off')
    fig.patch.set_facecolor('#f8f9fa')

    def svc_box(x, y, label, color=VERT, w=2.0, h=0.8):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.07",
                     facecolor=(*color, 0.12), edgecolor=color, lw=1.5))
        ax.text(x+w/2, y+h/2, label, ha='center', va='center',
                fontsize=8, fontweight='bold', color=color)
        return x+w/2, y+h

    def queue(x, y, label, color=(0.8,0.4,0.0)):
        ax.add_patch(FancyBboxPatch((x, y), 1.6, 0.55, boxstyle="round,pad=0.05",
                     facecolor=(*color, 0.15), edgecolor=color, lw=1.2, linestyle='--'))
        ax.text(x+0.8, y+0.275, label, ha='center', va='center', fontsize=6.5, color=color)

    def publish(x1, y1, x2, y2, label='publish', color=(0.8,0.4,0.0)):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle='->', color=color, lw=1.3))
        ax.text((x1+x2)/2 + 0.1, (y1+y2)/2 + 0.1, label,
                ha='center', fontsize=6.5, color=color, style='italic')

    ax.add_patch(FancyBboxPatch((5.5, 3.3), 3.0, 1.4, boxstyle="round,pad=0.1",
                 facecolor='#fff7ed', edgecolor=(0.8,0.4,0.0), lw=2))
    ax.text(7.0, 4.3, 'RabbitMQ', ha='center', fontsize=10, fontweight='bold',
            color=(0.8,0.4,0.0))
    ax.text(7.0, 3.8, 'Exchange : mfh.events', ha='center', fontsize=8, color='#666666')

    queues_pos = [
        (5.6, 2.5, 'loan.created'),
        (7.3, 2.5, 'payment.pending'),
        (5.6, 1.7, 'account.opened'),
        (7.3, 1.7, 'mobile.confirmed'),
    ]
    for qx, qy, ql in queues_pos:
        queue(qx, qy, ql)

    publishers = [
        (0.5, 6.5, 'Auth Service', (0.5,0.2,0.6)),
        (0.5, 5.2, 'Loan Service', VERT),
        (0.5, 3.9, 'Account Service', ORANGE),
        (0.5, 2.6, 'Transaction Svc', BLEU),
    ]
    for px, py, pl, pc in publishers:
        svc_box(px, py, pl, pc)
        publish(px+2.0, py+0.4, 5.5, 3.7)

    consumers = [
        (11.5, 6.5, 'Notification Svc', ROUGE),
        (11.5, 5.2, 'Repayment Svc', (0.5,0.2,0.6)),
        (11.5, 3.9, 'Loan Service', VERT),
        (11.5, 2.6, 'Client Service', BLEU),
    ]
    for cx, cy, cl, cc in consumers:
        svc_box(cx, cy, cl, cc)
        publish(8.5, 3.7, cx, cy+0.4, 'consume', cc)

    ax.set_title('Communication Asynchrone — RabbitMQ Message Broker',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=10)
    return fig_to_stream(fig)

# ══════════════════════════════════════════════════════════════════════════════
# DIAGRAMME 13 : Flux de Déploiement Docker
# ══════════════════════════════════════════════════════════════════════════════
def draw_deployment():
    fig, ax = plt.subplots(figsize=(16, 10))
    ax.set_xlim(0, 16); ax.set_ylim(0, 10); ax.axis('off')
    fig.patch.set_facecolor('#f8f9fa')

    def container(x, y, w, h, name, port, color=VERT, icon=''):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.06",
                     facecolor=(*color, 0.1), edgecolor=color, lw=1.5))
        ax.text(x+w/2, y+h*0.72, f'{icon} {name}' if icon else name,
                ha='center', va='center', fontsize=7.5, fontweight='bold', color=color)
        ax.text(x+w/2, y+h*0.28, port, ha='center', va='center',
                fontsize=6.5, color='#555555')

    # Hote Docker
    ax.add_patch(FancyBboxPatch((0.2, 0.3), 15.6, 9.2, boxstyle="round,pad=0.1",
                 facecolor='#f0f4ff', edgecolor='#3b82f6', lw=2, linestyle='--'))
    ax.text(8.0, 9.6, 'Machine Hote Linux — Docker Engine',
            ha='center', fontsize=11, fontweight='bold', color='#3b82f6')

    # Reseau Docker interne
    ax.add_patch(FancyBboxPatch((0.5, 0.5), 15.0, 8.8, boxstyle="round,pad=0.08",
                 facecolor='#f8fff8', edgecolor=VERT_L, lw=1.5, linestyle=':'))
    ax.text(8.0, 9.2, 'Reseau Docker : mfh-network (bridge)',
            ha='center', fontsize=9, color=VERT_L, style='italic')

    # Infrastructure (ligne 1)
    container(0.8, 0.8, 2.4, 1.0, 'PostgreSQL', '5433:5432', (0.3,0.5,0.8), 'DB')
    container(3.4, 0.8, 2.4, 1.0, 'RabbitMQ', '5672+15672', (0.8,0.4,0.0), 'MQ')
    container(6.0, 0.8, 2.4, 1.0, 'Redis', '6379:6379', (0.7,0.1,0.1), 'Cache')
    container(8.6, 0.8, 2.4, 1.0, 'Config Svc', '8000:8000', ORANGE, 'Cfg')
    container(11.2, 0.8, 3.5, 1.0, 'Eureka Registry', '8761:8761', ORANGE, 'Reg')

    # Ligne 2 — services metier 1/2
    svcs_row1 = [
        ('Auth Svc',    ':8080', 0.8,  2.2),
        ('Client Svc',  ':8081', 3.4,  2.2),
        ('Account Svc', ':8082', 6.0,  2.2),
        ('Loan Svc',    ':8083', 8.6,  2.2),
        ('Repayment',   ':8084', 11.2, 2.2),
    ]
    for name, port, x, y in svcs_row1:
        container(x, y, 2.4, 1.0, name, port, VERT)

    # Ligne 3 — services metier 2/2
    svcs_row2 = [
        ('Reporting',   ':8085', 0.8,  3.5),
        ('Agency Svc',  ':8086', 3.4,  3.5),
        ('Config.MFI',  ':8087', 6.0,  3.5),
        ('Transaction', ':8088', 8.6,  3.5),
        ('Notif. Svc',  ':8089', 11.2, 3.5),
    ]
    for name, port, x, y in svcs_row2:
        container(x, y, 2.4, 1.0, name, port, VERT)

    # Gateway & Frontend
    container(1.5, 5.0, 5.0, 1.4, 'API Gateway', ':8091', BLEU, 'GW')
    container(9.0, 5.0, 5.0, 1.4, 'Frontend React', '3000:80 (Nginx)', VERT_L, 'UI')

    # Volumes
    container(0.8, 7.0, 5.0, 1.0, 'Volumes persistants',
              'postgres_data | rabbitmq_data | redis_data', (0.4,0.4,0.4), 'VOL')

    # Labels couches
    for label, y_pos, col in [
        ('Infrastructure', 1.3, (0.3,0.5,0.8)),
        ('Services Metier', 2.7, VERT),
        ('Gateway / Frontend', 5.7, BLEU),
    ]:
        ax.text(15.4, y_pos, label, ha='right', va='center',
                fontsize=7, color=col, fontweight='bold', style='italic')

    ax.set_title('Diagramme de Deploiement Docker Compose — MicrofinanceHub',
                 fontsize=13, fontweight='bold', color='#064e3b', pad=12)
    return fig_to_stream(fig)

# ══════════════════════════════════════════════════════════════════════════════
# UTILITAIRES WORD
# ══════════════════════════════════════════════════════════════════════════════
def add_heading(doc, text, level, color_hex='#064e3b'):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex.lstrip('#'))
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text).font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)

def add_image(doc, stream, width_inches=6.5, caption=''):
    doc.add_picture(stream, width=Inches(width_inches))
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 6'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.cell(r+1, c)
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# GÉNÉRATION DU DOCUMENT WORD
# ══════════════════════════════════════════════════════════════════════════════
def make_report():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # ── PAGE DE GARDE ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('RAPPORT DE PROJET')
    run.bold = True; run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x06, 0x4e, 0x3b)

    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('MicrofinanceHub')
    r2.bold = True; r2.font.size = Pt(22)
    r2.font.color.rgb = RGBColor(0x06, 0x4e, 0x3b)

    doc.add_paragraph()
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run('Plateforme Microservices de Gestion de Microfinance').font.size = Pt(14)

    doc.add_paragraph(); doc.add_paragraph()
    p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(f'Date : {datetime.date.today().strftime("%d %B %Y")}').font.size = Pt(12)

    doc.add_paragraph()
    p5 = doc.add_paragraph(); p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.add_run('Architecture : Microservices Spring Boot | Frontend : React').font.size = Pt(11)

    doc.add_page_break()

    # ── TABLE DES MATIERES ─────────────────────────────────────────────────
    add_heading(doc, 'Table des Matieres', 1)
    toc = [
        ('1.',      'Introduction et Contexte'),
        ('2.',      'Presentation du Projet'),
        ('3.',      'Architecture Globale'),
        ('4.',      'Architecture Technique — Stack'),
        ('5.',      'Diagrammes UML'),
        ('  5.1.',  "Cas d'Utilisation"),
        ('  5.2.',  'Sequence — Authentification JWT'),
        ('  5.3.',  'Sequence — Demande de Pret'),
        ('  5.4.',  "Sequence — Transaction (depot)"),
        ('  5.5.',  'Classes Globales'),
        ('  5.6.',  'Architecture par Couche (service typique)'),
        ('  5.7.',  'Etats — Compte Bancaire'),
        ('  5.8.',  'Etats — Pret'),
        ('  5.9.',  'Etats — Paiement/Remboursement'),
        ('6.',      'Communication Inter-Services (RabbitMQ)'),
        ('7.',      'Modele de Donnees'),
        ('8.',      'Securite et Authentification'),
        ('9.',      'Services Developpes'),
        ('10.',     'Interface Utilisateur'),
        ('11.',     'Deploiement et Infrastructure'),
        ('12.',     'Conclusion et Perspectives'),
    ]
    for num, title in toc:
        p = doc.add_paragraph()
        p.add_run(f'{num}  {title}').font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ── 1. INTRODUCTION ────────────────────────────────────────────────────
    add_heading(doc, '1. Introduction et Contexte', 1)
    add_para(doc,
        "Le secteur de la microfinance joue un role crucial dans l'inclusion financiere, "
        "notamment dans les economies emergentes. Les institutions de microfinance (IMF) gerent "
        "quotidiennement une grande diversite d'operations : gestion des clients, des comptes "
        "bancaires, des prets, des remboursements et des paiements mobiles. Ce volume d'operations "
        "necessite un systeme d'information robuste, scalable et securise.")
    add_para(doc,
        "C'est dans ce contexte que s'inscrit MicrofinanceHub : une plateforme complete de gestion "
        "de microfinance construite sur une architecture microservices moderne. Le projet repond aux "
        "exigences de performance, de maintenabilite et d'evolutivite des systemes d'information "
        "contemporains.")

    add_heading(doc, 'Objectifs du Projet', 2)
    for b in [
        "Digitaliser l'ensemble du cycle de vie d'une institution de microfinance",
        "Garantir la separation des responsabilites via une architecture microservices",
        "Assurer une securite robuste basee sur JWT et RBAC (4 roles : Admin, Directeur, Agent, Client)",
        "Integrer les paiements mobiles (Mobile Money via CamPay)",
        "Offrir des tableaux de bord adaptes a chaque role utilisateur",
        "Deployer l'ensemble via Docker Compose pour une portabilite maximale",
    ]:
        add_bullet(doc, b)

    doc.add_page_break()

    # ── 2. PRESENTATION ────────────────────────────────────────────────────
    add_heading(doc, '2. Presentation du Projet', 1)
    add_para(doc,
        "MicrofinanceHub est une application web full-stack composee de 12 microservices "
        "backend et d'un frontend React. Elle couvre l'ensemble du perimetre fonctionnel "
        "d'une IMF moderne : de la creation du profil client jusqu'au reporting financier avance.")

    add_heading(doc, '2.1 Roles Utilisateurs', 2)
    add_table(doc,
        ['Role', 'Responsabilites', 'Acces'],
        [
            ('Administrateur', 'Gestion globale : utilisateurs, agences, validations',
             'Toutes sections, toutes agences'),
            ("Directeur d'Agence", 'Supervision agence : clients, agents, validations, prets',
             'Donnees de son agence uniquement'),
            ('Agent de Terrain', 'Operations quotidiennes : clients, comptes, guichet',
             'Ses clients assignes'),
            ('Client', 'Consultation, paiements, suivi des prets',
             'Son profil uniquement'),
        ]
    )

    add_heading(doc, '2.2 Fonctionnalites Cles', 2)
    for f in [
        "Gestion complete des clients (KYC, documents, score credit)",
        "Ouverture et gestion de comptes bancaires (Courant, Epargne, Micro-epargne)",
        "Cycle de pret complet : demande -> eligibilite -> approbation -> decaissement -> remboursement",
        "Calcul automatique du plan d'amortissement",
        "Paiements Multi-canal : Especes, Mobile Money (CamPay), Virement, Cheque",
        "Notifications temps reel : Email (Thymeleaf) + SMS (Twilio)",
        "Transactions financieres : Depot, Retrait, Virement inter-comptes",
        "Rapports et tableaux de bord par agence et globaux",
        "Gestion multi-agences avec hierarchie Directeur/Agent",
        "Stockage documents KYC sur Google Drive",
    ]:
        add_bullet(doc, f)

    doc.add_page_break()

    # ── 3. ARCHITECTURE GLOBALE ────────────────────────────────────────────
    add_heading(doc, '3. Architecture Globale', 1)
    add_para(doc,
        "Le systeme adopte une architecture microservices ou chaque service est independant, "
        "deployable separement et responsable d'un domaine metier precis. Les services "
        "communiquent via REST (synchrone) et RabbitMQ (asynchrone).")

    print("  -> Generation Architecture Globale...")
    add_image(doc, draw_architecture_globale(), 6.5,
              'Figure 1 — Architecture Globale MicrofinanceHub')

    add_heading(doc, '3.1 Inventaire des Services', 2)
    add_table(doc,
        ['Service', 'Port', 'Role', 'Base de donnees'],
        [
            ('config-service',        '8000', 'Serveur de configuration centralise (Spring Cloud Config)', '—'),
            ('registry-service',      '8761', 'Service Discovery (Eureka)', '—'),
            ('api-gateway',           '8091', "Point d'entree unique, routage, CORS", '—'),
            ('auth-service',          '8080', 'Authentification, JWT, sessions Redis', 'microfinancehub_auth'),
            ('client-service',        '8081', 'Profils clients, documents KYC, score credit', 'microfinancehub_client'),
            ('account-service',       '8082', 'Comptes bancaires, operations', 'microfinancehub_account'),
            ('loan-service',          '8083', 'Prets, amortissement, eligibilite', 'microfinancehub_loan'),
            ('repayment-service',     '8084', 'Remboursements, validation paiements', 'microfinancehub_repayment'),
            ('reporting-service',     '8085', 'Rapports Excel/PDF, statistiques', 'microfinancehub_reporting'),
            ('agency-service',        '8086', 'Agences, agents, directeurs', 'microfinancehub_agency'),
            ('configuration-service', '8087', 'Parametrage metier MFI', 'microfinancehub_configuration'),
            ('transaction-service',   '8088', 'Depots, retraits, virements, CamPay', 'microfinancehub_transaction'),
            ('notification-service',  '8089', 'Emails Thymeleaf, SMS Twilio, WebSocket', 'microfinancehub_notification'),
            ('app-frontend',          '3000', 'SPA React (Nginx en production)', '—'),
        ]
    )

    doc.add_page_break()

    # ── 4. STACK TECHNIQUE ─────────────────────────────────────────────────
    add_heading(doc, '4. Architecture Technique — Stack Technologique', 1)

    print("  -> Generation Stack Technique...")
    add_image(doc, draw_stack_technique(), 6.2,
              'Figure 2 — Stack Technologique par couche')

    add_heading(doc, '4.1 Technologies Backend', 2)
    add_table(doc,
        ['Technologie', 'Version', 'Usage'],
        [
            ('Spring Boot',      '3.4.4',    'Framework principal de chaque microservice'),
            ('Spring Cloud',     '2024.0.0', 'Config Server, Eureka, OpenFeign, Gateway'),
            ('Spring Security',  '6.x',      'Authentification, autorisation par role (RBAC)'),
            ('Spring Data JPA',  '3.x',      'ORM avec Hibernate, acces base de donnees'),
            ('JJWT',             '0.12.6',   'Generation et validation des tokens JWT'),
            ('OpenFeign',        '4.x',      'Appels REST inter-services declaratifs'),
            ('RabbitMQ (AMQP)', '3.13',      'Messagerie asynchrone entre services'),
            ('PostgreSQL',       '16',       'Base de donnees relationnelle (une par service)'),
            ('Redis',            '7',        'Cache sessions, rate limiting'),
            ('Lombok',           '1.18',     'Reduction du code boilerplate Java'),
            ('Thymeleaf',        '3.x',      'Templates HTML pour emails'),
        ]
    )

    add_heading(doc, '4.2 Technologies Frontend', 2)
    add_table(doc,
        ['Technologie', 'Usage'],
        [
            ('React 18',        'Framework SPA, gestion etat local et hooks'),
            ('Axios',           'Client HTTP pour appels REST, intercepteurs JWT'),
            ('Tailwind CSS',    'Framework CSS utilitaire, design system'),
            ('React Router v6', 'Navigation SPA, routes protegees par role'),
            ('Lucide React',    "Bibliotheque d'icones SVG"),
            ('LocalStorage',    'Persistance session (token JWT + profil utilisateur)'),
        ]
    )

    doc.add_page_break()

    # ── 5. DIAGRAMMES UML ──────────────────────────────────────────────────
    add_heading(doc, '5. Diagrammes UML', 1)

    # 5.1 Use Case
    add_heading(doc, "5.1 Diagramme de Cas d'Utilisation", 2)
    add_para(doc,
        "Le diagramme suivant represente les interactions entre les 4 acteurs principaux "
        "du systeme et les fonctionnalites auxquelles ils ont acces.")
    print("  -> Generation Use Case...")
    add_image(doc, draw_use_case(), 6.5,
              "Figure 3 — Diagramme de Cas d'Utilisation")

    doc.add_page_break()

    # 5.2 Sequence Auth
    add_heading(doc, '5.2 Diagramme de Sequence — Authentification JWT', 2)
    add_para(doc,
        "Ce diagramme illustre le flux complet d'authentification : depuis la saisie des "
        "identifiants jusqu'au stockage du JWT et a la redirection vers le dashboard approprie. "
        "Le token contient : l'email, le role, les noms et l'agencyId de l'utilisateur.")
    print("  -> Generation Sequence Authentification...")
    add_image(doc, draw_sequence_auth(), 6.2,
              'Figure 4 — Sequence Authentification JWT')

    # 5.3 Sequence Pret
    add_heading(doc, '5.3 Diagramme de Sequence — Demande de Pret', 2)
    add_para(doc,
        "Ce diagramme montre le flux complet de traitement d'une demande de pret, "
        "incluant la verification JWT, la consultation du profil client, le calcul "
        "de l'amortissement et la notification asynchrone via RabbitMQ.")
    print("  -> Generation Sequence Pret...")
    add_image(doc, draw_sequence_pret(), 6.5,
              'Figure 5 — Sequence Demande de Pret')

    doc.add_page_break()

    # 5.4 Sequence Transaction (NOUVEAU)
    add_heading(doc, "5.4 Diagramme de Sequence — Transaction (Depot d'Argent)", 2)
    add_para(doc,
        "Ce diagramme illustre le flux d'un depot d'argent effectue par un agent ou directeur. "
        "La requete transite par l'API Gateway (validation JWT), puis le Transaction Service "
        "credits le compte via un appel interne Feign vers l'Account Service, enregistre "
        "la transaction et publie un evenement asynchrone vers le Notification Service.")
    print("  -> Generation Sequence Transaction...")
    add_image(doc, draw_sequence_transaction(), 6.5,
              "Figure 6 — Sequence Depot d'Argent (Transaction Service)")

    doc.add_page_break()

    # 5.5 Classes — diagramme global
    add_heading(doc, '5.5 Diagramme de Classes Global', 2)
    add_para(doc,
        "Ce diagramme regroupe la totalite des entites metier de tous les microservices. "
        "Chaque zone coloree correspond a un service avec sa propre base de donnees PostgreSQL "
        "(pattern Database per Service). Les references cross-services utilisent des IDs "
        "simples (String/Long) plutot que des jointures JPA.")
    print("  -> Generation Diagramme de Classes Global...")
    add_image(doc, draw_classes(), 6.5,
              'Figure 7 — Diagramme de Classes Global (10 microservices)')

    doc.add_page_break()

    # 5.6 Diagrammes par service
    add_heading(doc, '5.6 Diagrammes de Classes par Service', 2)
    add_para(doc,
        "Les diagrammes suivants detaillent les entites (avec tous les attributs) "
        "et les enumerations de chaque microservice.")

    svc_diagrams = [
        ('5.6.1 auth-service (:8080)',
         draw_classes_auth_service,
         'Figure 8a — auth-service : User, Role, Privilege, RefreshToken, UserSession, AuditLog, PasswordResetToken'),
        ('5.6.2 client-service (:8081)',
         draw_classes_client_service,
         'Figure 8b — client-service : Client, Document'),
        ('5.6.3 account-service (:8082)',
         draw_classes_account_service,
         'Figure 8c — account-service : Compte'),
        ('5.6.4 agency-service (:8086)',
         draw_classes_agency_service,
         'Figure 8d — agency-service : Agency, AgentAssignment, AgencyStats, AgencyConfiguration'),
        ('5.6.5 loan-service (:8083)',
         draw_classes_loan_service,
         'Figure 8e — loan-service : LoanProduct, LoanApplication, Loan, AmortizationSchedule, Schedule'),
        ('5.6.6 repayment-service (:8084)',
         draw_classes_repayment_service,
         'Figure 8f — repayment-service : Repayment, Payment, Schedule, Penalty'),
        ('5.6.7 transaction-service (:8088)',
         draw_classes_transaction_service,
         'Figure 8g — transaction-service : Transaction'),
        ('5.6.8 notification-service (:8089)',
         draw_classes_notification_service,
         'Figure 8h — notification-service : Notification, NotificationTemplate'),
        ('5.6.9 reporting-service (:8085)',
         draw_classes_reporting_service,
         'Figure 8i — reporting-service : Report, Kpi'),
        ('5.6.10 configuration-service (:8087)',
         draw_classes_configuration_service,
         'Figure 8j — configuration-service : MicrofinanceConfiguration, AccountCategory, AccountTypeConfiguration'),
    ]
    for (title, fn, caption) in svc_diagrams:
        add_heading(doc, title, 3)
        print(f'  -> Generation {title}...')
        add_image(doc, fn(), 6.5, caption)
        doc.add_page_break()

    doc.add_page_break()

    # 5.7 Etats Compte (NOUVEAU)
    add_heading(doc, '5.7 Diagramme d\'Etats — Compte Bancaire', 2)
    add_para(doc,
        "Un compte bancaire suit un cycle de vie strict : cree en EN_ATTENTE_VALIDATION, "
        "il doit etre approuve par un directeur ou admin (-> ACTIF). Il peut ensuite etre "
        "bloque, debloque ou ferme. Un rejet lors de la validation le place en etat REJETE.")
    print("  -> Generation Etats Compte...")
    add_image(doc, draw_state_compte(), 6.2,
              'Figure 9 — Diagramme d\'Etats : Compte Bancaire (StatutCompte)')

    # 5.8 Etats Pret (NOUVEAU)
    add_heading(doc, "5.8 Diagramme d'Etats — Pret", 2)
    add_para(doc,
        "Le cycle de vie d'un pret commence a l'etat PENDING apres soumission de la demande. "
        "Un admin ou directeur peut l'approuver (-> APPROVED) ou le rejeter (-> REJECTED). "
        "Apres decaissement, il passe ACTIVE, puis COMPLETED a remboursement total ou "
        "DEFAULTED en cas de defaut de paiement.")
    print("  -> Generation Etats Pret...")
    add_image(doc, draw_state_loan(), 6.2,
              "Figure 10 — Diagramme d'Etats : Pret (LoanStatus)")

    doc.add_page_break()

    # 5.9 Etats Paiement (NOUVEAU)
    add_heading(doc, "5.9 Diagramme d'Etats — Paiement/Remboursement", 2)
    add_para(doc,
        "Un paiement enregistre par un agent ou directeur passe en PENDING_VALIDATION "
        "et attend la validation de l'admin. Un paiement direct client ou un webhook "
        "CamPay peuvent passer directement en COMPLETED. En cas de probleme, le statut "
        "passe en FAILED.")
    print("  -> Generation Etats Paiement...")
    add_image(doc, draw_state_payment(), 6.2,
              "Figure 11 — Diagramme d'Etats : Paiement/Remboursement (PaymentStatus)")

    doc.add_page_break()

    # ── 6. RABBITMQ ────────────────────────────────────────────────────────
    add_heading(doc, '6. Communication Inter-Services (RabbitMQ)', 1)
    add_para(doc,
        "La communication asynchrone via RabbitMQ permet de decoupler les services et "
        "d'assurer la resilience du systeme. Lorsqu'un service publie un evenement, "
        "il n'attend pas la reponse des consommateurs, ce qui ameliore les performances "
        "et la disponibilite globale.")

    print("  -> Generation Diagramme RabbitMQ...")
    add_image(doc, draw_rabbitmq(), 6.2,
              'Figure 12 — Communication Asynchrone via RabbitMQ')

    add_heading(doc, '6.1 Evenements Principaux', 2)
    add_table(doc,
        ['Evenement', 'Producteur', 'Consommateur', 'Declencheur'],
        [
            ('user.login',        'Auth Service',      'Client Service',    'Connexion utilisateur'),
            ('loan.created',      'Loan Service',      'Notification Svc',  'Nouvelle demande pret'),
            ('loan.approved',     'Loan Service',      'Notification Svc',  'Pret approuve'),
            ('payment.pending',   'Repayment Service', 'Notification Svc',  'Paiement en attente'),
            ('payment.validated', 'Repayment Service', 'Loan Service',      'Paiement valide -> MAJ echeance'),
            ('account.opened',    'Account Service',   'Notification Svc',  'Compte ouvert'),
            ('mobile.confirmed',  'Transaction Svc',   'Repayment Service', 'CamPay webhook recu'),
            ('payment.received',  'Repayment Service', 'Notification Svc',  'Paiement CamPay confirme'),
        ]
    )

    doc.add_page_break()

    # ── 7. MODELE DE DONNEES ────────────────────────────────────────────────
    add_heading(doc, '7. Modele de Donnees', 1)
    add_para(doc,
        "Le projet applique le pattern 'Database per Service' : chaque microservice "
        "possede sa propre base de donnees PostgreSQL, garantissant l'independance "
        "et l'isolation des donnees. Les jointures inter-services se font via des "
        "appels REST (Feign) et non par des cles etrangeres croisees.")

    add_heading(doc, '7.1 Bases de Donnees', 2)
    add_table(doc,
        ['Base de donnees', 'Service', 'Tables principales', 'Port hote'],
        [
            ('microfinancehub_auth',          'auth-service',          'users, roles, refresh_tokens, user_sessions', '5433'),
            ('microfinancehub_client',        'client-service',        'clients, documents', '5433'),
            ('microfinancehub_account',       'account-service',       'compte, compte_events', '5433'),
            ('microfinancehub_loan',          'loan-service',          'loans, loan_applications, loan_schedules', '5433'),
            ('microfinancehub_repayment',     'repayment-service',     'payments, repayments, schedules', '5433'),
            ('microfinancehub_reporting',     'reporting-service',     'reports, report_data', '5433'),
            ('microfinancehub_agency',        'agency-service',        'agencies, agent_assignments, agency_stats', '5433'),
            ('microfinancehub_configuration', 'configuration-service', 'mfi_configs, loan_products', '5433'),
            ('microfinancehub_transaction',   'transaction-service',   'transactions, campay_webhooks', '5433'),
            ('microfinancehub_notification',  'notification-service',  'notifications, notification_logs', '5433'),
        ]
    )

    add_heading(doc, '7.2 Enums et Types Metier', 2)
    add_table(doc,
        ['Enum', 'Service', 'Valeurs'],
        [
            ('UserRoleType',     'Auth',        'ADMIN, DIRECTEUR_AGENCE, AGENT, CLIENT'),
            ('ClientStatus',     'Client',      'ACTIVE, INACTIVE, PENDING, SUSPENDED'),
            ('StatutCompte',     'Account',     'ACTIF, INACTIF, SUSPENDU, BLOQUE, FERME, EN_ATTENTE_VALIDATION, REJETE'),
            ('TypeCompte',       'Account',     'COURANT, EPARGNE, MICRO_EPARGNE, DEPOT_A_TERME, CREDIT'),
            ('LoanStatus',       'Loan',        'PENDING, APPROVED, REJECTED, ACTIVE, COMPLETED, DEFAULT'),
            ('PaymentStatus',    'Repayment',   'PENDING, PENDING_VALIDATION, COMPLETED, FAILED'),
            ('PaymentMethod',    'Repayment',   'CASH, MOBILE_MONEY, BANK_TRANSFER, CHECK'),
            ('TypeNotification', 'Notification','DEMANDE_PRET, APPROBATION_PRET, CONFIRMATION_REMB, ALERTE_RETARD, ...'),
        ]
    )

    doc.add_page_break()

    # ── 8. SECURITE ─────────────────────────────────────────────────────────
    add_heading(doc, '8. Securite et Authentification', 1)
    add_para(doc,
        "La securite est implementee a plusieurs niveaux pour garantir la confidentialite "
        "et l'integrite des donnees financieres.")

    add_heading(doc, '8.1 JWT et RBAC', 2)
    add_para(doc,
        "Chaque requete authentifiee inclut un Bearer Token JWT signe avec une cle secrete "
        "partagee (HMAC-SHA256). Le token contient les claims suivants :")
    for c in [
        "sub : email de l'utilisateur (identifiant unique)",
        "role : role utilisateur (ADMIN / DIRECTEUR_AGENCE / AGENT / CLIENT)",
        "firstName / lastName : nom complet",
        "agencyId / agencyCode : identifiant de l'agence (pour agents et directeurs)",
        "exp : date d'expiration (24h par defaut)",
    ]:
        add_bullet(doc, c)

    add_heading(doc, "8.2 Controle d'Acces par Endpoint", 2)
    add_table(doc,
        ['Endpoint', 'Roles Autorises'],
        [
            ('POST /api/loans/approve',                 'ADMIN, DIRECTEUR_AGENCE'),
            ('GET /api/clients/by-agent/{email}',       'ADMIN, DIRECTEUR_AGENCE'),
            ('POST /api/transactions/depot/**',         'CLIENT, AGENT, ADMIN, DIRECTEUR_AGENCE'),
            ('PATCH /api/comptes/{id}/statut',          'ADMIN, DIRECTEUR_AGENCE'),
            ('GET /api/agency/my-clients',              'DIRECTEUR_AGENCE, ADMIN'),
            ('POST /api/clients/internal/by-agent-emails', 'Public (interne)'),
            ('DELETE /api/comptes/{id}',                'ADMIN uniquement'),
        ]
    )

    add_heading(doc, '8.3 Mesures de Securite Additionnelles', 2)
    for s in [
        "BCrypt (force 12) pour le hachage des mots de passe",
        "Rate Limiting (via filtre personnalise) pour limiter les tentatives de connexion",
        "Sessions Redis avec TTL pour l'invalidation centralisee des sessions",
        "CORS configure par service (origines autorisees explicitement)",
        "Validation des entrees avec Bean Validation (@Valid, @NotNull, etc.)",
        "Tokens de service internes (ADMIN JWT) pour les appels Feign inter-services",
        "Secrets externalises via variables d'environnement Docker",
    ]:
        add_bullet(doc, s)

    doc.add_page_break()

    # ── 9. SERVICES ─────────────────────────────────────────────────────────
    add_heading(doc, '9. Services Developpes — Detail Fonctionnel', 1)

    services_detail = [
        ('Auth Service (8080)',
         "Gere l'inscription, la connexion, et la gestion des tokens. "
         "Supporte 4 roles (ADMIN, DIRECTEUR_AGENCE, AGENT, CLIENT). "
         "Les tokens JWT incluent le role, les noms et l'agencyId. "
         "Les sessions sont stockees dans Redis avec invalidation a la deconnexion. "
         "Inclut la recuperation de mot de passe par email.",
         ['POST /auth/login — Authentification + generation JWT',
          'POST /auth/register — Inscription client',
          'POST /auth/agent/create — Creation agent (admin)',
          'GET /auth/me — Profil utilisateur connecte',
          'PUT /internal/users/{id}/agency — Mise a jour agence (interne)']),

        ('Client Service (8081)',
         "Gere les profils clients de l'institution. Chaque client est associe a un agent "
         "(createdBy) et optionnellement a une agence (agencyId). Le service calcule et "
         "met a jour le score de credit. Les documents KYC sont stockes sur Google Drive.",
         ['POST /api/clients — Creation client',
          'GET /api/clients/by-agent/{email} — Clients d\'un agent',
          'GET /api/clients/by-agency/{id} — Clients d\'une agence',
          'POST /api/clients/internal/by-agent-emails — Interne (agency-service)',
          'PATCH /api/clients/{id}/status — Changement statut']),

        ('Account Service (8082)',
         "Gere les comptes bancaires. Un compte passe par l'etat EN_ATTENTE_VALIDATION "
         "avant d'etre active par un directeur ou admin. Le service expose des endpoints "
         "internes pour les operations de credit/debit appeles par transaction-service.",
         ['POST /api/comptes — Ouverture de compte',
          'GET /api/comptes/en-attente-validation — File de validation',
          'PATCH /api/comptes/{id}/statut — Activation/Rejet (directeur/admin)',
          'POST /api/comptes/{id}/crediter — Interne (transaction-service)',
          'GET /api/comptes/client/{id} — Comptes d\'un client']),

        ('Loan Service (8083)',
         "Gere le cycle complet du pret : demande, eligibilite, amortissement, "
         "approbation, decaissement, suivi des echeances. Calcule automatiquement "
         "les plans d'amortissement (methode des interets degressifs).",
         ['POST /api/loans/apply — Demande de pret',
          'GET /api/loans/eligibility/{clientId} — Verification eligibilite',
          "GET /api/loans/{id}/amortization — Plan d'amortissement",
          'POST /api/loans/approval/{id}/approve — Approbation',
          'POST /api/loans/by-clients — Prets pour liste de clients']),

        ('Repayment Service (8084)',
         "Gere les remboursements de prets. Les agents et directeurs enregistrent "
         "des paiements en PENDING_VALIDATION ; l'admin les valide. CamPay declenche "
         "les confirmations via RabbitMQ (webhook asynchrone).",
         ['POST /api/repayments/pay/record — Enregistrement (agent/directeur)',
          'POST /api/repayments/pay/client — Paiement direct client',
          'GET /api/repayments/pending — Paiements en attente (admin)',
          'POST /api/repayments/{id}/validate — Validation (admin)',
          'POST /api/repayments/stats/by-clients — Stats par agence']),

        ('Agency Service (8086)',
         "Gere les agences, l'assignation des agents et directeurs, et expose "
         "les donnees agregees des clients par agence. La methode getAgencyClientsWithAccounts "
         "recupere les clients via les emails des agents assignes (endpoint interne client-service).",
         ['GET /api/agency/my-agency — Agence du directeur connecte',
          "GET /api/agency/my-clients — Clients de l'agence (via agents)",
          'GET /api/agency/my-agents — Agents (actifs + inactifs)',
          'PATCH /api/agency/agents/{id}/toggle-status — Activer/Desactiver agent',
          'POST /api/agency/agents/assign — Assigner un agent']),

        ('Transaction Service (8088)',
         "Gere toutes les operations financieres : depots, retraits, virements "
         "inter-comptes et integration CamPay (Mobile Money). Les transactions "
         "modifient les soldes en appelant account-service via Feign.",
         ['POST /api/transactions/depot/{compteId} — Depot',
          'POST /api/transactions/retrait/{compteId} — Retrait',
          'POST /api/transactions/virement/{compteId} — Virement',
          'POST /api/campay/webhook — Webhook CamPay',
          'GET /api/transactions/compte/{compteId} — Historique']),

        ('Notification Service (8089)',
         "Envoie des notifications aux clients et agents via Email (templates Thymeleaf) "
         "et SMS (Twilio). Les notifications sont declenchees de maniere asynchrone "
         "via RabbitMQ. Supporte les notifications en temps reel via WebSocket.",
         ['Emails : pret approuve, compte ouvert, paiement recu',
          "SMS : alertes critiques, rappels d'echeance",
          'Templates HTML Thymeleaf avec mise en page professionnelle',
          'WebSocket pour notifications push en temps reel']),
    ]

    for svc_name, svc_desc, svc_endpoints in services_detail:
        add_heading(doc, svc_name, 2)
        add_para(doc, svc_desc)
        add_para(doc, 'Endpoints principaux :', bold=True)
        for ep in svc_endpoints:
            add_bullet(doc, ep)
        doc.add_paragraph()

    doc.add_page_break()

    # ── 10. INTERFACE UTILISATEUR ───────────────────────────────────────────
    add_heading(doc, '10. Interface Utilisateur', 1)
    add_para(doc,
        "Le frontend React est organise en espaces distincts selon le role de l'utilisateur. "
        "Chaque role dispose d'un layout personnalise avec une navigation adaptee a ses "
        "responsabilites.")

    add_heading(doc, '10.1 Espaces Utilisateurs', 2)
    add_table(doc,
        ['Role', 'Route', 'Pages disponibles'],
        [
            ('Admin',     '/admin/**',     'Dashboard, Clients, Prets, Comptes, Validations, Transactions, Remboursements, Rapports, Notifications, Agents, Directeurs, Agences, Parametres'),
            ('Directeur', '/directeur/**', 'Dashboard, Clients (CRUD), Agents (toggle), Validations, Comptes, Prets, Remboursements, Rapports, Guichet'),
            ('Agent',     '/agent/**',     'Dashboard, Mes Clients, Guichet, Demandes Pret, Remboursements, Mon Profil'),
            ('Client',    '/client/**',    'Mon Espace, Mes Comptes, Mes Prets, Transactions, Remboursements, Simulateur'),
        ]
    )

    add_heading(doc, '10.2 Fonctionnalites Cles du Directeur', 2)
    for f in [
        "Vue consolidee des clients de son agence (via agents assignes)",
        "Creation de nouveaux clients avec assignation d'agence",
        "Modification et changement de statut des clients",
        "Toggle actif/inactif des agents de son agence",
        "Validation des demandes de compte (EN_ATTENTE -> ACTIF ou REJETE)",
        "Enregistrement de paiements de remboursement (-> PENDING_VALIDATION pour l'admin)",
        "Operations guichet : Depot, Retrait, Virement pour les clients de l'agence",
        "Rapports financiers filtres par agence",
    ]:
        add_bullet(doc, f)

    doc.add_page_break()

    # ── 11. DEPLOIEMENT ─────────────────────────────────────────────────────
    add_heading(doc, '11. Deploiement et Infrastructure', 1)
    add_para(doc,
        "L'ensemble du systeme est conteneurise via Docker et orchestre par Docker Compose. "
        "Chaque service dispose de son propre Dockerfile multi-etapes (build Maven + image JRE legere). "
        "Le frontend utilise Nginx pour servir les fichiers React buildes.")

    print("  -> Generation Diagramme Deploiement...")
    add_image(doc, draw_deployment(), 6.5,
              'Figure 13 — Architecture de Deploiement Docker Compose')

    add_heading(doc, "11.1 Variables d'Environnement (Docker Secrets)", 2)
    add_table(doc,
        ['Variable', 'Usage'],
        [
            ('JWT_SECRET',                        'Cle HMAC-SHA256 partagee (256 bits) pour signer les JWT'),
            ('DB_USERNAME / DB_PASSWORD',         'Credentials PostgreSQL'),
            ('RABBITMQ_USERNAME / PASSWORD',      'Credentials RabbitMQ'),
            ('MAIL_HOST / PORT / USERNAME / PASSWORD', 'Serveur SMTP pour emails'),
            ('CAMPAY_USERNAME / PASSWORD',        'Credentials API CamPay (Mobile Money)'),
            ('TWILIO_ACCOUNT_SID / AUTH_TOKEN',   'Credentials Twilio pour SMS'),
            ('GOOGLE_DRIVE_ROOT_FOLDER_ID',       'Dossier Google Drive pour KYC documents'),
            ('ADMIN_EMAIL / ADMIN_PASSWORD',      'Compte admin cree au demarrage'),
        ]
    )

    add_heading(doc, '11.2 Ordre de Demarrage', 2)
    add_para(doc, "Les services respectent un ordre de demarrage via les conditions healthcheck de Docker Compose :")
    for s in [
        "1. Infrastructure : postgres, rabbitmq, redis",
        "2. Spring Cloud Core : config-service -> registry-service",
        "3. API Gateway",
        "4. Services metier (parallele) : auth, client, account, agency, loan, repayment, reporting, configuration, transaction, notification",
        "5. Frontend : app-frontend",
    ]:
        add_bullet(doc, s)

    doc.add_page_break()

    # ── 12. CONCLUSION ──────────────────────────────────────────────────────
    add_heading(doc, '12. Conclusion et Perspectives', 1)
    add_para(doc,
        "MicrofinanceHub demontre la faisabilite de la construction d'une application "
        "financiere complete selon les principes des architectures microservices modernes. "
        "Le projet integre les meilleures pratiques du developpement logiciel contemporain : "
        "separation des responsabilites, communication asynchrone, securite par defaut, "
        "et deploiement conteneurise.")

    add_heading(doc, 'Points Forts', 2)
    for s in [
        "Architecture scalable : chaque service peut etre mis a l'echelle independamment",
        "Securite multicouche : JWT RBAC + BCrypt + Redis sessions + Rate Limiting",
        "Resilience : communication asynchrone RabbitMQ, gestion des erreurs Feign",
        "Experience utilisateur : interfaces adaptees a chaque role, actions contextuelles",
        "Integration tiers : CamPay (Mobile Money), Twilio (SMS), Google Drive (KYC)",
        "Observabilite : logs structures, Spring Actuator health checks",
        "Portabilite : deploiement complet en une commande Docker Compose",
    ]:
        add_bullet(doc, s)

    add_heading(doc, "Perspectives d'Evolution", 2)
    for e in [
        "Migration vers Kubernetes pour l'orchestration en production",
        "Centralisation des logs avec ELK Stack (Elasticsearch, Logstash, Kibana)",
        "Mise en place de Keycloak pour la gestion centralisee des identites (SSO)",
        "Circuit Breaker (Resilience4j) pour la tolerance aux pannes entre services",
        "Application mobile native (React Native / Flutter)",
        "Intelligence artificielle pour la notation de credit (ML credit scoring)",
        "Audit trail complet et conformite reglementaire",
        "API publique documentee (OpenAPI 3.0) pour partenaires tiers",
    ]:
        add_bullet(doc, e)

    doc.add_paragraph(); doc.add_paragraph()
    p_final = doc.add_paragraph()
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_final = p_final.add_run('— MicrofinanceHub — Rapport de Presentation —')
    run_final.font.size = Pt(10)
    run_final.font.italic = True
    run_final.font.color.rgb = RGBColor(0x06, 0x4e, 0x3b)

    output_path = '/home/axelle-mbadi/IdeaProjects/MicrofinanceHub/Rapport_MicrofinanceHub.docx'
    doc.save(output_path)
    print(f"\nRapport genere : {output_path}")
    return output_path

if __name__ == '__main__':
    print("Generation du rapport MicrofinanceHub...")
    path = make_report()
    print(f"Fichier : {path}")
